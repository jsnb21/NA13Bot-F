"""
NA13Bot-F - Main Flask Application
====================================
Core Flask application entry point for the multi-tenant restaurant AI chatbot system.
Handles all HTTP routing, session management, file uploads, and coordination between
frontend, database, and AI services.

Key Responsibilities:
  - Flask app initialization and configuration
  - Authentication routes (signup, login, OTP verification)
  - Admin dashboard for restaurant owners
  - Chatbot interface for customers
  - File upload management (logos, training data, menu items)
  - Training data processing and indexing
  - Database initialization
  - Session and cookie management
  - AI training data integration

Main Page Routes:
  - GET / - Redirect to dashboard/login
  - GET /login - Login page
  - GET /signup - Signup page
  - GET /dashboard - Admin dashboard
  - GET /chatbot - Customer chatbot interface
  - GET /settings - Settings page
  - GET /ai-training - AI training management
  - GET /menu - Menu management
  - GET /orders - Order history
  - GET /reports - Analytics and reports

API Routes:
  - POST /upload-training-file - Training file upload
  - POST /delete-training-file - Remove training file
  - GET /get-trained-brands - List trained restaurants
  - GET /brand/image/<kind>/<restaurant_id> - Serve logo/avatar from database
  - POST /save-settings - Save restaurant configuration (includes logo/avatar DB upload)
  - POST /save-menu-description - Update menu text
  - POST /api/* - Chatbot API endpoints (see chatbot/routes.py)

Features:
  - Multi-tenant restaurant support with isolated data
  - OTP-based authentication with 5-minute TTL
  - Flask-Turbo for seamless real-time navigation
  - Secure file handling with validation and sanitization
  - Training data manifest management
  - PDF processing support (via PyPDF)
  - CSV/JSON/DOCX file parsing
  - Database-backed configuration storage
  - Session-based access control

Configuration:
  - FLASK_SECRET: Session encryption key (from env or default)
  - ALLOWED_EXT: Logo file extensions (png, jpg, jpeg, gif, svg)
  - TRAINING_ALLOWED_EXT: Training file extensions (txt, pdf, docx, json, csv)
  - MAX_TRAINING_FILE_MB: Maximum training file size (50MB)
  - OTP_TTL_SECONDS: OTP validity (300 seconds / 5 minutes)

Key Functions:
  - allowed_file(): Validate logo file extensions
  - training_allowed_file(): Validate training file extensions
  - get_training_dir(): Get restaurant training directory path
  - get_training_manifest_path(): Get manifest.json path
  - load_training_manifest(): Load training file metadata
  - save_training_manifest(): Save training file metadata
  - save_training_upload_bytes(): Save training file to disk
  - _store_brand_image_upload(): Store logo/avatar binary data to database
  - _migrate_static_brand_image(): Migrate legacy /static/uploads images to database
  - extract_pdf_text(): Parse PDF files
  - extract_docx_text(): Parse DOCX files
  - extract_csv_text(): Parse CSV files

Dependencies:
  - Flask: Web framework
  - Flask-Turbo: Real-time page updates
  - Jinja2: Template rendering
  - PyPDF: PDF processing
  - psycopg: PostgreSQL driver
  - google.genai: Gemini AI API
"""

from flask import Flask, render_template, request, redirect, url_for, jsonify, session, Response, make_response
from flask_turbo import Turbo
from tools import (
    save_config,
    load_config,
    add_user,
    verify_user,
    user_exists,
    get_user,
    update_user_meta,
    normalize_menu_item_name,
)
from config import init_db, get_connection, get_db_schema
import os
from werkzeug.utils import secure_filename
from pathlib import Path
import json
import csv
import io
import re
import mimetypes
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None
try:
    from docx import Document
except Exception:
    Document = None
import time
import secrets
import uuid
import hashlib
import socket
from datetime import datetime, timezone, timedelta
from urllib.parse import urlparse, urlunparse
from chatbot.routes import chatbot_bp
from chatbot.training import (
    build_training_context,
    build_training_chunks,
    get_training_dir as training_get_dir,
    get_training_manifest_path as training_get_manifest_path,
    get_training_history_path as training_get_history_path,
    load_training_manifest as training_load_manifest,
    save_training_manifest as training_save_manifest,
    load_training_history as training_load_history,
    save_training_history as training_save_history,
)
import google.genai as genai
import base64
import psycopg
import smtplib
import ssl
from email.message import EmailMessage
import qrcode
from io import BytesIO

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv(dotenv_path=Path(__file__).resolve().parent / '.env', override=True)

# after app is created, before routes
init_db()

# simple session secret
app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET', 'dev-secret')

# Initialize Turbo for seamless page navigation
turbo = Turbo(app)

# Add custom Jinja2 filter for regex operations
@app.template_filter('regex_findall')
def regex_findall_filter(text, pattern):
    """Find all matches of a regex pattern in text"""
    import re
    if not text:
        return []
    return re.findall(pattern, str(text))

app.register_blueprint(chatbot_bp)
from functools import wraps
from flask import flash

# allowed logo extensions
ALLOWED_EXT = {'png', 'jpg', 'jpeg', 'gif', 'svg', 'webp'}
TRAINING_ALLOWED_EXT = {'txt', 'pdf', 'docx', 'json', 'csv'}
MAX_TRAINING_FILE_MB = 50
OTP_TTL_SECONDS = 300

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXT


def training_allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in TRAINING_ALLOWED_EXT


BRAND_IMAGE_COLUMNS = {
    'logo': ('logo_data', 'logo_mime'),
    'chatbot_avatar': ('chatbot_avatar_data', 'chatbot_avatar_mime')
}


def _detect_image_mime(image_bytes: bytes, fallback: str = 'application/octet-stream') -> str:
    """Best-effort image mime detection to avoid serving images as text/plain/octet-stream."""
    if not image_bytes:
        return fallback

    head = bytes(image_bytes[:32])
    if head.startswith(b'\x89PNG\r\n\x1a\n'):
        return 'image/png'
    if head.startswith(b'\xff\xd8\xff'):
        return 'image/jpeg'
    if head.startswith(b'GIF87a') or head.startswith(b'GIF89a'):
        return 'image/gif'
    if head.startswith(b'RIFF') and b'WEBP' in bytes(image_bytes[8:16]):
        return 'image/webp'

    # Basic SVG detection (handles optional BOM/whitespace)
    probe = bytes(image_bytes[:512]).lstrip(b'\xef\xbb\xbf\x00\t\r\n ')
    if probe.startswith(b'<?xml') or probe.startswith(b'<svg'):
        return 'image/svg+xml'

    return fallback


def _normalize_image_mime(mime_type: str, image_bytes: bytes) -> str:
    mime = (mime_type or '').strip().lower()
    if mime.startswith('image/'):
        return mime
    return _detect_image_mime(image_bytes, fallback='application/octet-stream')


def _get_brand_image_url(image_kind: str, restaurant_id: str) -> str:
    version = f"{int(time.time() * 1000)}-{uuid.uuid4().hex[:8]}"
    return f"/brand/image/{image_kind}/{restaurant_id}?v={version}"


def _store_brand_image_bytes(restaurant_id: str, image_kind: str, image_bytes: bytes, mime_type: str) -> str:
    if not restaurant_id or image_kind not in BRAND_IMAGE_COLUMNS:
        raise ValueError('Invalid brand image target')
    if not image_bytes:
        raise ValueError('Image data is empty')

    data_col, mime_col = BRAND_IMAGE_COLUMNS[image_kind]
    normalized_mime = _normalize_image_mime(mime_type, image_bytes)
    schema = get_db_schema()

    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                psycopg.sql.SQL(
                    """
                    INSERT INTO {}.brand_settings (restaurant_id, {}, {})
                    VALUES (%s, %s, %s)
                    ON CONFLICT (restaurant_id) DO UPDATE
                    SET {} = EXCLUDED.{}, {} = EXCLUDED.{}, updated_at = now()
                    """
                ).format(
                    psycopg.sql.Identifier(schema),
                    psycopg.sql.Identifier(data_col),
                    psycopg.sql.Identifier(mime_col),
                    psycopg.sql.Identifier(data_col),
                    psycopg.sql.Identifier(data_col),
                    psycopg.sql.Identifier(mime_col),
                    psycopg.sql.Identifier(mime_col)
                ),
                [restaurant_id, image_bytes, normalized_mime]
            )

    return _get_brand_image_url(image_kind, restaurant_id)


def _store_brand_image_upload(restaurant_id: str, image_kind: str, upload_file) -> str:
    safe_name = secure_filename(upload_file.filename or '')
    ext = Path(safe_name).suffix.lower()
    if not ext or ext.lstrip('.') not in ALLOWED_EXT:
        raise ValueError('Unsupported file type')

    image_bytes = upload_file.read()
    mime_type = upload_file.content_type
    if not mime_type or mime_type == 'application/octet-stream':
        guessed_mime, _ = mimetypes.guess_type(safe_name)
        mime_type = guessed_mime or 'application/octet-stream'

    mime_type = _normalize_image_mime(mime_type, image_bytes)

    return _store_brand_image_bytes(restaurant_id, image_kind, image_bytes, mime_type)


def _migrate_static_brand_image(restaurant_id: str, image_kind: str, image_url: str) -> str:
    if not image_url or not image_url.startswith('/static/uploads/'):
        return image_url

    project_root = Path(__file__).parent.resolve()
    candidate = (project_root / image_url.lstrip('/')).resolve()

    try:
        candidate.relative_to(project_root)
    except Exception:
        return image_url

    if not candidate.is_file():
        return image_url

    ext = candidate.suffix.lower().lstrip('.')
    if ext not in ALLOWED_EXT:
        return image_url

    image_bytes = candidate.read_bytes()
    guessed_mime, _ = mimetypes.guess_type(candidate.name)
    return _store_brand_image_bytes(restaurant_id, image_kind, image_bytes, guessed_mime or 'application/octet-stream')


def get_training_dir(restaurant_id: str):
    return training_get_dir(restaurant_id)


def get_training_manifest_path(restaurant_id: str):
    return training_get_manifest_path(restaurant_id)


def get_training_history_path(restaurant_id: str):
    return training_get_history_path(restaurant_id)


def load_training_manifest(restaurant_id: str):
    return training_load_manifest(restaurant_id)


def save_training_manifest(restaurant_id: str, entries):
    training_save_manifest(restaurant_id, entries)


def load_training_history(restaurant_id: str):
    return training_load_history(restaurant_id)


def save_training_history(restaurant_id: str, entries):
    training_save_history(restaurant_id, entries)


def add_training_history_entry(restaurant_id: str, entry: dict, max_entries: int = 200):
    entries = load_training_history(restaurant_id)
    entries.append(entry)
    if len(entries) > max_entries:
        entries = entries[-max_entries:]
    save_training_history(restaurant_id, entries)


def save_training_upload_bytes(restaurant_id: str, filename: str, data_bytes: bytes):
    if not filename:
        return None
    if not training_allowed_file(filename):
        return None
    training_dir = get_training_dir(restaurant_id)
    safe_name = secure_filename(filename)
    ext = Path(safe_name).suffix.lower()
    stored_name = f"{uuid.uuid4().hex}{ext}"
    dest = training_dir / stored_name
    dest.write_bytes(data_bytes)

    entry = {
        'id': uuid.uuid4().hex,
        'original_name': safe_name,
        'stored_name': stored_name,
        'uploaded_at': datetime.now(timezone.utc).isoformat(),
        'status': 'ready',
        'size_bytes': len(data_bytes or b'')
    }
    entries = load_training_manifest(restaurant_id)
    entries.append(entry)
    save_training_manifest(restaurant_id, entries)
    return entry


def _canonical_size_label(label: str) -> str:
    token = (label or '').strip().lower()
    mapping = {
        's': 'Small',
        'sm': 'Small',
        'small': 'Small',
        'm': 'Medium',
        'md': 'Medium',
        'med': 'Medium',
        'medium': 'Medium',
        'l': 'Large',
        'lg': 'Large',
        'large': 'Large',
    }
    return mapping.get(token, '')


_CURRENCY_CODE_PATTERN = re.compile(
    r"\b(?:php|usd|aud|cad|eur|gbp|jpy|sgd|inr|hkd|cny|rmb|myr|thb|vnd|idr|krw|chf|nzd)\.?\b",
    re.IGNORECASE,
)
_CURRENCY_SYMBOL_PATTERN = re.compile(r"(?:A\$|C\$|S\$|HK\$|US\$|CA\$|AU\$|NZ\$|R\$|[$€£¥₹₱]|[\u20A0-\u20CF])")
_CURRENCY_SYMBOL_MAP = {
    'PHP': '₱',
    'USD': '$',
    'EUR': '€',
    'GBP': '£',
    'JPY': '¥',
    'AUD': 'A$',
    'CAD': 'C$',
    'SGD': 'S$',
    'INR': '₹',
    'HKD': 'HK$',
    'NZD': 'NZ$',
    'BRL': 'R$',
}
_CURRENCY_CODE_SYMBOL_HINTS = {
    'PHP': ('₱',),
    'USD': ('US$', '$'),
    'EUR': ('€',),
    'GBP': ('£',),
    'JPY': ('¥',),
    'AUD': ('A$',),
    'CAD': ('C$', 'CA$'),
    'SGD': ('S$',),
    'INR': ('₹',),
    'HKD': ('HK$',),
    'NZD': ('NZ$',),
    'BRL': ('R$',),
}


def _detect_currency_from_text(content: str):
    text = str(content or '')
    if not text.strip():
        return None

    scores = {}

    def _bump(code: str, amount: int):
        if not code or amount <= 0:
            return
        scores[code] = scores.get(code, 0) + amount

    upper_text = text.upper()
    for code in _CURRENCY_CODE_SYMBOL_HINTS.keys():
        code_hits = len(re.findall(rf"\\b{re.escape(code)}\\.?\\b", upper_text))
        if code_hits:
            _bump(code, code_hits * 3)

    symbol_hints = []
    for code, symbols in _CURRENCY_CODE_SYMBOL_HINTS.items():
        for symbol in symbols:
            symbol_hints.append((symbol, code))
    symbol_hints.sort(key=lambda pair: len(pair[0]), reverse=True)

    for symbol, code in symbol_hints:
        hits = len(re.findall(re.escape(symbol), text))
        if not hits:
            continue
        weight = 1 if symbol == '$' else 2
        _bump(code, hits * weight)

    if not scores:
        return None

    ranked = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    if len(ranked) > 1 and ranked[0][1] == ranked[1][1]:
        return None

    detected_code, confidence = ranked[0]
    return {
        'code': detected_code,
        'symbol': _CURRENCY_SYMBOL_MAP.get(detected_code, ''),
        'confidence': confidence,
    }


def _build_currency_mismatch_warning(configured_code: str, configured_symbol: str, detected_currency):
    if not detected_currency:
        return None

    configured_code_value = (configured_code or '').strip().upper()
    configured_symbol_value = (configured_symbol or '').strip()
    detected_code = (detected_currency.get('code') or '').strip().upper()
    detected_symbol = (detected_currency.get('symbol') or '').strip()
    if not configured_code_value or not detected_code:
        return None

    if configured_code_value == detected_code:
        return None

    configured_label = configured_code_value
    if configured_symbol_value:
        configured_label = f"{configured_code_value} ({configured_symbol_value})"

    detected_label = detected_code
    if detected_symbol:
        detected_label = f"{detected_code} ({detected_symbol})"

    return (
        f"This PDF appears to use {detected_label}, but your web-app currency is set to {configured_label}. "
        "Please review imported prices carefully, as they may become inaccurate."
    )


def _strip_currency_tokens(price_text: str) -> str:
    text = str(price_text or '').strip()
    if not text:
        return ''
    text = _CURRENCY_CODE_PATTERN.sub('', text)
    text = _CURRENCY_SYMBOL_PATTERN.sub('', text)
    text = re.sub(r"\s{2,}", ' ', text)
    return text.strip(" \t:-")


def _parse_price_value(price_text: str):
    clean = _strip_currency_tokens(price_text)
    if not clean:
        return None
    match = re.search(r"\d[\d,]*(?:\.\d{1,2})?", clean)
    if not match:
        return None
    try:
        return float(match.group(0).replace(',', ''))
    except Exception:
        return None


def _parse_size_variant_line(line: str):
    clean = (line or '').strip()
    if not clean:
        return None

    variant_pattern = re.compile(
        r"\b(?P<label>small|medium|large|sm|md|lg|s|m|l)\b\s*[:=\-]?\s*(?P<price>(?:[$₱]|php\.?|usd\.?|aud\.?|cad\.?|eur\.?)?\s*\d+(?:\.\d{1,2})?)",
        re.IGNORECASE
    )
    matches = list(variant_pattern.finditer(clean))
    if len(matches) < 2:
        return None

    name = clean[:matches[0].start()].strip(" -:\t|,")
    if not name:
        return None

    options = []
    seen = set()
    for match in matches:
        label = _canonical_size_label(match.group('label'))
        price = _strip_currency_tokens(match.group('price'))
        if not label or not price or label in seen:
            continue
        seen.add(label)
        options.append((label, price))

    if len(options) < 2:
        return None

    option_text = ', '.join(f"{label}={price}" for label, price in options)
    numeric_prices = [val for _, price in options if (val := _parse_price_value(price)) is not None]
    base_price = str(int(min(numeric_prices))) if numeric_prices and min(numeric_prices).is_integer() else (str(min(numeric_prices)) if numeric_prices else options[0][1])

    return {
        'name': name,
        'description': f"Options: {option_text}",
        'price': base_price,
        'category': 'Uncategorized',
        'status': 'Live'
    }


def _merge_small_medium_large_variants(items):
    if not items:
        return []

    variant_name_pattern = re.compile(r"^(?P<base>.+?)\s*(?:[-(]\s*)?(?P<label>small|medium|large|sm|md|lg|s|m|l)\s*\)?$", re.IGNORECASE)
    merged = []
    grouped = {}

    for item in items:
        if not isinstance(item, dict):
            continue
        name = (item.get('name') or '').strip()
        if not name:
            continue

        description = (item.get('description') or '').strip()
        if description.lower().startswith('options:'):
            merged.append(item)
            continue

        match = variant_name_pattern.match(name)
        if not match:
            merged.append(item)
            continue

        label = _canonical_size_label(match.group('label'))
        base_name = (match.group('base') or '').strip(" -:\t|,")
        price = _strip_currency_tokens(item.get('price'))
        if not label or not base_name or not price:
            merged.append(item)
            continue

        category = (item.get('category') or 'Uncategorized').strip() or 'Uncategorized'
        key = (base_name.lower(), category.lower())
        if key not in grouped:
            grouped[key] = {
                'item': {
                    'name': base_name,
                    'description': '',
                    'price': price,
                    'category': category,
                    'status': (item.get('status') or 'Live').strip() or 'Live'
                },
                'options': [],
                'seen_labels': set(),
                'prices': []
            }
        group = grouped[key]
        if label not in group['seen_labels']:
            group['options'].append((label, price))
            group['seen_labels'].add(label)
        numeric = _parse_price_value(price)
        if numeric is not None:
            group['prices'].append(numeric)
        if description and not group['item']['description']:
            group['item']['description'] = description

    for key, group in grouped.items():
        options = group['options']
        if len(options) < 2:
            # If we only collected one size label, preserve original naming behavior by skipping merge.
            label, price = options[0]
            suffix_item = dict(group['item'])
            suffix_item['name'] = f"{group['item']['name']} {label}"
            suffix_item['price'] = price
            merged.append(suffix_item)
            continue

        order = {'Small': 0, 'Medium': 1, 'Large': 2}
        options.sort(key=lambda row: order.get(row[0], 99))
        option_text = ', '.join(f"{label}={price}" for label, price in options)
        merged_item = group['item']
        merged_item['description'] = f"Options: {option_text}"
        if group['prices']:
            minimum = min(group['prices'])
            merged_item['price'] = str(int(minimum)) if minimum.is_integer() else str(minimum)
        merged.append(merged_item)

    return merged


def _strip_options_block(description: str) -> str:
    text = (description or '').strip()
    if not text:
        return ''
    return re.sub(r'\boptions\s*:\s*.*$', '', text, flags=re.IGNORECASE).strip(' -|,;')


def _build_variant_options_from_form(form_data):
    options = []

    # Preferred format: customizable label + price pairs.
    default_labels = ['Small', 'Medium', 'Large']
    for idx in range(1, 4):
        label = (form_data.get(f'variant_label_{idx}') or '').strip()
        value = _strip_currency_tokens(form_data.get(f'variant_price_{idx}'))
        if not value:
            continue
        if not label:
            label = default_labels[idx - 1]
        options.append((label, value))

    # Backward compatibility with older fixed field names.
    if not options:
        for label, key in (('Small', 'variant_small'), ('Medium', 'variant_medium'), ('Large', 'variant_large')):
            value = _strip_currency_tokens(form_data.get(key))
            if value:
                options.append((label, value))

    if not options:
        return None

    option_text = ', '.join(f"{label}={price}" for label, price in options)
    numeric_prices = [val for _, price in options if (val := _parse_price_value(price)) is not None]
    if numeric_prices:
        minimum = min(numeric_prices)
        price_value = str(int(minimum)) if minimum.is_integer() else str(minimum)
    else:
        price_value = options[0][1]

    return {
        'description_options': f"Options: {option_text}",
        'price': price_value
    }


def _apply_variant_form_values(description: str, price: str, form_data):
    variant_data = _build_variant_options_from_form(form_data)
    if not variant_data:
        return description, _strip_currency_tokens(price)

    base_description = _strip_options_block(description)
    if base_description:
        full_description = f"{base_description}. {variant_data['description_options']}"
    else:
        full_description = variant_data['description_options']
    return full_description, variant_data['price']


def parse_menu_txt(content: str):
    if not content:
        return []
    text = content.replace('\r\n', '\n').replace('\r', '\n')

    items = []
    heading_matches = list(re.finditer(r"\b([A-Z][A-Z &]{2,})\b(?=\s+NAME:)", text))
    headings = [(m.start(), m.group(1).strip()) for m in heading_matches]

    pattern = re.compile(
        r"NAME:\s*(.*?)\s*\|\s*PRICE:\s*(.*?)\s*\|\s*DESCRIPTION:\s*(.*?)(?=\s+(?:[A-Z][A-Z &]{2,}\s+)?NAME:|$)",
        re.DOTALL
    )

    for match in pattern.finditer(text):
        name = (match.group(1) or '').strip()
        if not name:
            continue
        price = _strip_currency_tokens(match.group(2))
        description = (match.group(3) or '').strip()
        category = 'Uncategorized'
        for pos, heading in reversed(headings):
            if pos < match.start():
                category = heading.title()
                break
        items.append({
            'name': name,
            'description': description,
            'price': price,
            'category': category,
            'status': 'Live'
        })

    if items:
        return items

    price_pattern = re.compile(r"(\$\s*\d+(?:\.\d{1,2})?|\d+(?:\.\d{1,2})?\s*(?:usd|php|php\.|aud|cad|eur)?)", re.IGNORECASE)
    for line in text.split('\n'):
        clean = line.strip()
        if not clean:
            continue
        variant_item = _parse_size_variant_line(clean)
        if variant_item:
            items.append(variant_item)
            continue
        match = price_pattern.search(clean)
        if not match:
            continue
        price = _strip_currency_tokens(match.group(1))
        name_part = clean[:match.start()].strip(" -:\t")
        desc_part = clean[match.end():].strip(" -:\t")
        if not name_part:
            continue
        items.append({
            'name': name_part,
            'description': desc_part,
            'price': price,
            'category': 'Uncategorized',
            'status': 'Live'
        })
    if items:
        return _merge_small_medium_large_variants(items)

    rows = []
    reader = csv.reader(io.StringIO(content))
    for row in reader:
        if not row:
            continue
        cleaned = [cell.strip() for cell in row]
        if not any(cleaned):
            continue
        rows.append(cleaned)

    if not rows:
        return []

    first = [cell.lower() for cell in rows[0]]
    if 'name' in first and ('price' in first or 'description' in first):
        rows = rows[1:]

    for row in rows:
        name = row[0].strip() if len(row) > 0 else ''
        if not name:
            continue
        description = row[1].strip() if len(row) > 1 else ''
        price = _strip_currency_tokens(row[2]) if len(row) > 2 else ''
        category = row[3].strip() if len(row) > 3 else 'Uncategorized'
        status = row[4].strip() if len(row) > 4 else 'Live'
        items.append({
            'name': name,
            'description': description,
            'price': price,
            'category': category or 'Uncategorized',
            'status': status or 'Live'
        })
    return _merge_small_medium_large_variants(items)


def _extract_known_categories(menu_items):
    known = []
    for item in menu_items or []:
        cat = (item.get('category') or '').strip()
        if cat and cat not in known:
            known.append(cat)
    return known


def infer_menu_category(name: str, description: str = '', known_categories=None) -> str:
    text = f"{name or ''} {description or ''}".lower()
    known_categories = known_categories or []

    # Prefer matching already-used restaurant categories if any of their terms appear.
    known_scores = []
    for category in known_categories:
        clean = (category or '').strip()
        if not clean or clean.lower() == 'uncategorized':
            continue
        tokens = [tok for tok in re.split(r'[^a-z0-9]+', clean.lower()) if len(tok) >= 3]
        score = sum(1 for token in tokens if token in text)
        if score > 0:
            known_scores.append((score, clean))
    if known_scores:
        known_scores.sort(key=lambda row: (-row[0], row[1]))
        return known_scores[0][1]

    keyword_map = [
        ('Drinks/Beverages', [
            'drink', 'beverage', 'coffee', 'latte', 'espresso', 'mocha', 'cappuccino',
            'tea', 'juice', 'soda', 'cola', 'shake', 'smoothie', 'milk tea', 'boba',
            'frappe', 'cocktail', 'mocktail', 'wine', 'beer'
        ]),
        ('Desserts', [
            'dessert', 'cake', 'ice cream', 'gelato', 'brownie', 'cookie', 'pastry',
            'donut', 'doughnut', 'pie', 'pudding', 'sweet'
        ]),
        ('Appetizers', [
            'appetizer', 'starter', 'side', 'fries', 'nachos', 'wings', 'dumpling',
            'spring roll', 'salad', 'soup', 'sampler'
        ]),
        ('Main Course', [
            'main', 'entree', 'rice', 'pasta', 'noodle', 'burger', 'pizza', 'steak',
            'chicken', 'beef', 'pork', 'seafood', 'fish', 'adobo', 'sisig', 'curry'
        ])
    ]

    for default_category, keywords in keyword_map:
        if any(keyword in text for keyword in keywords):
            return default_category

    return 'Uncategorized'


def _format_price_value(value):
    try:
        numeric = float(value)
    except Exception:
        return ''
    if numeric.is_integer():
        return str(int(numeric))
    return f"{numeric:.2f}".rstrip('0').rstrip('.')


def _extract_variant_options(description: str):
    text = (description or '').strip()
    if not text:
        return []

    match = re.search(r'\boptions\s*:\s*(.*)$', text, flags=re.IGNORECASE)
    if not match:
        return []

    variants = []
    seen = set()
    option_text = match.group(1)
    for token in option_text.split(','):
        part = token.strip()
        if '=' not in part:
            continue
        label, raw_price = part.split('=', 1)
        clean_label = (label or '').strip()
        clean_price = _strip_currency_tokens(raw_price)
        if not clean_label or not clean_price:
            continue
        key = clean_label.lower()
        if key in seen:
            continue
        seen.add(key)
        variants.append((clean_label, clean_price))
    return variants


def _extract_variants_from_freeform_text(text: str):
    source = (text or '').strip()
    if not source:
        return []

    variants = []
    seen = set()

    # Labeled format examples: "Small 120", "T=165", "1ea: 150", "SET - 280"
    labeled_pattern = re.compile(
        r"\b(?P<label>(?:small|medium|large|sm|md|lg|s|m|l|t|g|v|set|solo|regular|family|\d+\s*(?:ea|pcs?|pc)))\b\s*[:=\-]?\s*(?P<price>(?:[$₱]|php\.?|usd\.?|aud\.?|cad\.?|eur\.?)?\s*\d+(?:\.\d{1,2})?)",
        re.IGNORECASE
    )
    for match in labeled_pattern.finditer(source):
        label = (match.group('label') or '').strip()
        price = _strip_currency_tokens(match.group('price'))
        if not label or not price:
            continue
        key = label.lower()
        if key in seen:
            continue
        seen.add(key)
        variants.append((label, price))

    if len(variants) >= 2:
        return variants

    # Unlabeled numeric format examples: "165 180", "120 / 140 / 160"
    numbers = re.findall(r"(?:[$₱]|php\.?|usd\.?|aud\.?|cad\.?|eur\.?)?\s*\d+(?:\.\d{1,2})?", source, flags=re.IGNORECASE)
    alpha_chars = re.sub(r"[^A-Za-z]+", "", source)
    if len(numbers) >= 2 and len(alpha_chars) <= 3:
        default_labels = ['Small', 'Medium', 'Large', 'XL', 'XXL']
        parsed = []
        for idx, raw_price in enumerate(numbers[:5]):
            label = default_labels[idx] if idx < len(default_labels) else f"Option {idx + 1}"
            parsed.append((label, _strip_currency_tokens(raw_price)))
        return parsed

    return []


def _normalize_menu_status(status_value: str, force_draft: bool = False) -> str:
    status_map = {
        'live': 'Live',
        'draft': 'Draft',
        'hidden': 'Hidden'
    }
    normalized = status_map.get((status_value or '').strip().lower(), 'Live')
    if force_draft and normalized == 'Live':
        return 'Draft'
    return normalized


def _normalize_imported_menu_item(item: dict, known_categories):
    if not isinstance(item, dict):
        return None, False

    name = (item.get('name') or '').strip()
    if not name:
        return None, False

    description = (item.get('description') or '').strip()
    category = (item.get('category') or '').strip()
    if not category or category.lower() == 'uncategorized':
        category = infer_menu_category(name, description, known_categories)
    normalized_name = normalize_menu_item_name(name, category)
    if normalized_name:
        name = normalized_name

    base_price_text = _strip_currency_tokens(item.get('price'))
    base_numeric = _parse_price_value(base_price_text)

    variants = []
    for variant in item.get('variants') or []:
        if not isinstance(variant, dict):
            continue
        label = (variant.get('label') or '').strip()
        if not label:
            continue
        raw_price = variant.get('price')
        if isinstance(raw_price, dict):
            if raw_price.get('display'):
                price_text = str(raw_price.get('display')).strip()
            elif raw_price.get('amount') is not None:
                price_text = _format_price_value(raw_price.get('amount'))
            else:
                price_text = ''
        else:
            price_text = _strip_currency_tokens(raw_price)
        if not price_text:
            continue
        variants.append((label, price_text))

    if not variants:
        variants = _extract_variant_options(description)
    if not variants:
        variants = _extract_variants_from_freeform_text(description)

    variant_values = [val for _, p in variants if (val := _parse_price_value(p)) is not None]
    if base_numeric is None and isinstance(item.get('base_price'), dict):
        base_amount = item.get('base_price', {}).get('amount')
        if base_amount is not None:
            base_numeric = _parse_price_value(str(base_amount))

    if variant_values:
        minimum = min(variant_values)
        base_numeric = minimum if base_numeric is None else min(base_numeric, minimum)
        option_text = ', '.join(f"{label}={price}" for label, price in variants)
        description_core = _strip_options_block(description)
        description = f"{description_core}. Options: {option_text}".strip('. ') if description_core else f"Options: {option_text}"

    needs_review = bool(item.get('needs_review'))
    confidence_overall = None
    if isinstance(item.get('confidence'), dict):
        try:
            confidence_overall = float(item.get('confidence', {}).get('overall'))
        except Exception:
            confidence_overall = None
    if confidence_overall is not None and confidence_overall < 0.75:
        needs_review = True

    if base_numeric is None:
        base_price = ''
        needs_review = True
    else:
        base_price = _format_price_value(base_numeric)

    status = _normalize_menu_status(item.get('status'), force_draft=needs_review)

    normalized = {
        'name': name,
        'description': description,
        'price': base_price,
        'category': category,
        'status': status
    }
    if item.get('image_url'):
        normalized['image_url'] = item.get('image_url')
    if item.get('id'):
        normalized['id'] = item.get('id')

    return normalized, needs_review


def _extract_json_payload_from_text(text: str):
    raw = (text or '').strip()
    if not raw:
        return None

    try:
        return json.loads(raw)
    except Exception:
        pass

    # Fallback for model responses that wrap JSON with extra text.
    for opener, closer in (('{', '}'), ('[', ']')):
        start = raw.find(opener)
        end = raw.rfind(closer)
        if start == -1 or end == -1 or end <= start:
            continue
        try:
            return json.loads(raw[start:end + 1])
        except Exception:
            continue

    return None


def _generate_gemini_text(parts, *, system_instruction: str = '', response_mime_type: str = '', temperature: float = 0.2, max_output_tokens: int = 8000):
    api_key = get_google_api_key()
    if not api_key:
        return ''

    try:
        client = genai.Client(api_key=api_key)
        cfg_kwargs = {
            'temperature': temperature,
            'max_output_tokens': max_output_tokens,
        }
        if system_instruction:
            cfg_kwargs['system_instruction'] = system_instruction
        if response_mime_type:
            cfg_kwargs['response_mime_type'] = response_mime_type

        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[{"role": "user", "parts": parts}],
            config=genai.types.GenerateContentConfig(**cfg_kwargs)
        )
        return (response.text or '').strip()
    except Exception:
        return ''


def parse_training_text_with_ai(content: str, source_name: str = ''):
    text = (content or '').strip()
    if not text:
        return {}

    # Keep payload bounded to avoid oversized token usage on large manuals.
    sample = text[:18000]
    system_instruction = (
        "You analyze restaurant training documents and return compact JSON metadata only. "
        "Return object keys: document_type, categories, summary, key_points, faq. "
        "Rules: categories must be short labels, key_points max 8 items, faq max 8 items. "
        "Each faq item must include question and answer. "
        "Do not invent details that are not present in text."
    )

    prompt = (
        f"SOURCE: {source_name or 'training_file'}\\n"
        "Analyze this training content and produce JSON metadata for retrieval and admin review.\\n\\n"
        f"CONTENT:\\n{sample}"
    )

    raw = _generate_gemini_text(
        parts=[{"text": prompt}],
        system_instruction=system_instruction,
        response_mime_type='application/json',
        temperature=0.1,
        max_output_tokens=3000
    )
    payload = _extract_json_payload_from_text(raw)
    if not isinstance(payload, dict):
        return {}

    categories = payload.get('categories')
    if not isinstance(categories, list):
        categories = []
    categories = [str(c).strip() for c in categories if str(c).strip()][:8]

    key_points = payload.get('key_points')
    if not isinstance(key_points, list):
        key_points = []
    key_points = [str(p).strip() for p in key_points if str(p).strip()][:8]

    faq_rows = payload.get('faq')
    faq = []
    if isinstance(faq_rows, list):
        for row in faq_rows[:8]:
            if not isinstance(row, dict):
                continue
            q = str(row.get('question') or '').strip()
            a = str(row.get('answer') or '').strip()
            if q and a:
                faq.append({'question': q, 'answer': a})

    return {
        'document_type': str(payload.get('document_type') or '').strip(),
        'categories': categories,
        'summary': str(payload.get('summary') or '').strip(),
        'key_points': key_points,
        'faq': faq,
    }


def parse_menu_txt_with_ai(content: str):
    if not content:
        return []

    try:
        system_instruction = (
            "Extract ALL restaurant menu items from the provided text. DO NOT SKIP ANY ITEMS. "
            "Return JSON only: an array of objects with keys "
            "name, description, price, category, status, variants. "
            "Use empty string when a field is missing. "
            "Return prices as plain numbers without currency symbols or codes. "
            "If a section heading (e.g., APPETIZERS, ESPRESSO BEVERAGE) appears, use it as category. "
            "Default status to Live. "
            "DESCRIPTION RULES: The description should contain actual descriptive text about the item. "
            "Do NOT put standalone numbers (like calories, nutritional values) in the description. "
            "If there's no descriptive text, leave description as empty string. "
            "CRITICAL PRICING VARIANTS: If an item has multiple prices (sizes T/G/V, quantities 1ea/SET, portions, etc.), "
            "store them in variants as an array of objects [{label, price}] and preserve labels exactly. "
            "Put the lowest price in the price field. "
            "Create ONE item per menu item name, not separate items per variant. "
            "COMPLETENESS: Extract EVERY item in the text. Count them if needed to ensure none are missing."
        )
        text = _generate_gemini_text(
            parts=[{"text": content}],
            system_instruction=system_instruction,
            response_mime_type='application/json',
            temperature=0.2,
            max_output_tokens=8000
        )
        if not text:
            return []

        data = _extract_json_payload_from_text(text)

        if data is None:
            return []
        if isinstance(data, dict) and 'items' in data:
            data = data.get('items')
        if not isinstance(data, list):
            return []

        items = []
        known_categories = []
        for row in data:
            if not isinstance(row, dict):
                continue
            name = (row.get('name') or '').strip()
            if not name:
                continue
            description = (row.get('description') or '').strip()
            price = _strip_currency_tokens(row.get('price'))
            if not price and isinstance(row.get('base_price'), dict):
                amount = row.get('base_price', {}).get('amount')
                if amount is not None:
                    price = _format_price_value(amount)

            variants = []
            for variant in row.get('variants') or []:
                if not isinstance(variant, dict):
                    continue
                label = (variant.get('label') or '').strip()
                if not label:
                    continue
                raw_price = variant.get('price')
                if isinstance(raw_price, dict):
                    display = (raw_price.get('display') or '').strip()
                    amount = raw_price.get('amount')
                    variant_price = _strip_currency_tokens(display) or (_format_price_value(amount) if amount is not None else '')
                else:
                    variant_price = _strip_currency_tokens(raw_price)
                if variant_price:
                    variants.append((label, variant_price))

            if variants and 'options:' not in description.lower():
                option_text = ', '.join(f"{label}={val}" for label, val in variants)
                description = f"{description}. Options: {option_text}".strip('. ') if description else f"Options: {option_text}"
            elif not variants:
                variants = _extract_variants_from_freeform_text(description)
                if variants and 'options:' not in description.lower():
                    option_text = ', '.join(f"{label}={val}" for label, val in variants)
                    description = f"{description}. Options: {option_text}".strip('. ') if description else f"Options: {option_text}"

            needs_review = bool(row.get('needs_review'))
            if isinstance(row.get('confidence'), dict):
                try:
                    if float(row.get('confidence', {}).get('overall')) < 0.75:
                        needs_review = True
                except Exception:
                    pass

            category = (row.get('category') or '').strip() or infer_menu_category(name, description, known_categories)
            status = _normalize_menu_status(row.get('status'), force_draft=needs_review)
            if category and category not in known_categories and category.lower() != 'uncategorized':
                known_categories.append(category)
            items.append({
                'name': name,
                'description': description,
                'price': price,
                'category': category,
                'status': status
            })
        return _merge_small_medium_large_variants(items)
    except Exception:
        return []


def extract_image_text_with_ai(data_bytes: bytes, mime_type: str):
    if not data_bytes:
        return ''
    try:
        prompt = (
            "Extract ALL text from this restaurant menu image. This is CRITICAL - you must not skip any items.\n\n"
            "Instructions:\n"
            "- Extract EVERY SINGLE menu item visible in the image\n"
            "- Preserve the exact text as it appears\n"
            "- Keep each menu item on its own line\n"
            "- Preserve section headings (ESPRESSO, TEA, etc) in ALL CAPS\n"
            "- Include all prices and size labels (T, G, V, Small, Medium, Large, etc)\n"
            "- If items are in columns, read left to right, top to bottom\n"
            "- If text is small or faint, still extract it\n"
            "- Double-check you haven't missed any items\n\n"
            "Return ONLY the extracted text, nothing else. BE COMPLETE."
        )
        encoded = base64.b64encode(data_bytes).decode('ascii')
        text = _generate_gemini_text(
            parts=[
                {"text": prompt},
                {"inline_data": {"mime_type": mime_type, "data": encoded}}
            ],
            temperature=0.1,
            max_output_tokens=8000
        )
        if not text:
            print("Warning: Gemini API returned empty response for image extraction")
        return text
    except Exception as e:
        print(f"Error extracting text from image: {str(e)}")
        return ''


def extract_pdf_text(data_bytes: bytes):
    if not PdfReader:
        return ''
    try:
        reader = PdfReader(io.BytesIO(data_bytes))
    except Exception:
        return ''
    parts = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ''
        except Exception:
            text = ''
        if text:
            parts.append(text)
    return '\n'.join(parts)


def extract_docx_text(data_bytes: bytes):
    if not Document:
        return ''
    try:
        doc = Document(io.BytesIO(data_bytes))
    except Exception:
        return ''
    parts = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or '').strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts)

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('user'):
            return redirect(url_for('login', next=request.path))
        return f(*args, **kwargs)
    return decorated


def _get_superadmin_credentials():
    """Get superadmin credentials from environment variables only.
    
    SUPERADMIN_USER and SUPERADMIN_PASSWORD must be set in .env file.
    Falls back to defaults if not set, but config.json is never used for security.
    """
    default_user = 'admin'
    default_password = 'admin123'

    env_user = (os.environ.get('SUPERADMIN_USER') or '').strip()
    env_password = (os.environ.get('SUPERADMIN_PASSWORD') or '').strip()
    if env_user and env_password:
        return env_user, env_password

    return default_user, default_password


def _is_superadmin_authenticated():
    return bool(session.get('superadmin_authenticated'))


def _set_superadmin_authenticated(value: bool):
    if value:
        session['superadmin_authenticated'] = True
        session['is_superadmin'] = True
        return
    session.pop('superadmin_authenticated', None)
    session.pop('is_superadmin', None)


def _verify_superadmin_credentials(username: str, password: str):
    expected_user, expected_password = _get_superadmin_credentials()
    provided_user = (username or '').strip()
    provided_password = password or ''
    return (
        secrets.compare_digest(provided_user, expected_user)
        and secrets.compare_digest(provided_password, expected_password)
    )


def _verify_superadmin_password(password: str):
    _, expected_password = _get_superadmin_credentials()
    provided_password = password or ''
    return secrets.compare_digest(provided_password, expected_password)


def _superadmin_unauthorized_response():
    return jsonify({
        'ok': False,
        'message': 'Superadmin authentication required.'
    }), 401


def superadmin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if _is_superadmin_authenticated():
            return f(*args, **kwargs)
        return _superadmin_unauthorized_response()
    return decorated

def get_current_restaurant_id():
    rid = session.get('restaurant_id')
    if rid:
        email = session.get('user')
        ensure_brand_seed(rid, email)
        return rid
    email = session.get('user')
    if not email:
        return None
    user = get_user(email) or {}
    meta = user.get('meta') or {}
    rid = meta.get('restaurant_id')
    if not rid:
        rid = str(uuid.uuid4())
        update_user_meta(email, {'restaurant_id': rid})
    session['restaurant_id'] = rid
    ensure_brand_seed(rid, email)
    return rid


def ensure_brand_seed(restaurant_id: str, email: str):
    if not restaurant_id or not email:
        return
    current = load_config(restaurant_id)
    if current.get('establishment_name'):
        return
    user = get_user(email) or {}
    meta = user.get('meta') or {}
    seed = {
        'establishment_name': meta.get('establishment_name', ''),
        'logo_url': meta.get('logo_url', ''),
        'main_color': meta.get('main_color', ''),
        'sub_color': meta.get('sub_color', '')
    }
    if any(seed.values()):
        save_config(seed, restaurant_id)


def generate_otp_code():
    return f"{secrets.randbelow(1000000):06d}"


def set_otp(email: str, purpose: str):
    code = generate_otp_code()
    session['otp'] = {
        'email': email,
        'code': code,
        'purpose': purpose,
        'expires_at': time.time() + OTP_TTL_SECONDS
    }
    return code


def get_smtp_settings():
    host = os.environ.get('SMTP_HOST', '').strip()
    port_raw = os.environ.get('SMTP_PORT', '587').strip()
    try:
        port = int(port_raw)
    except ValueError:
        port = 587
    user = os.environ.get('SMTP_USER', '').strip()
    password = os.environ.get('SMTP_PASSWORD', '').strip()
    use_tls = os.environ.get('SMTP_USE_TLS', 'true').strip().lower() in ('1', 'true', 'yes', 'on')
    use_ssl = os.environ.get('SMTP_USE_SSL', 'false').strip().lower() in ('1', 'true', 'yes', 'on')
    from_email = os.environ.get('SMTP_FROM', '').strip() or user
    from_name = os.environ.get('SMTP_FROM_NAME', 'Resto AI').strip()
    return {
        'host': host,
        'port': port,
        'user': user,
        'password': password,
        'use_tls': use_tls,
        'use_ssl': use_ssl,
        'from_email': from_email,
        'from_name': from_name
    }


def should_show_otp_hint():
    return os.environ.get('OTP_DEBUG_SHOW', '').strip().lower() in ('1', 'true', 'yes', 'on')


def send_otp_email(to_email: str, code: str, purpose: str, cfg: dict = None):
    settings = get_smtp_settings()
    if not settings['host'] or not settings['from_email']:
        return False, 'SMTP is not configured.'

    cfg = cfg or {}
    brand = cfg.get('establishment_name') or 'Resto AI'
    ttl_minutes = max(1, int(OTP_TTL_SECONDS / 60))
    subject = f"Your {brand} verification code"
    intro = "Use this code to finish signing in" if purpose == 'login' else "Use this code to finish creating your account"

    body = (
        f"{intro}.\n\n"
        f"Verification code: {code}\n"
        f"This code expires in {ttl_minutes} minutes.\n\n"
        f"If you did not request this, you can ignore this email."
    )

    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = f"{settings['from_name']} <{settings['from_email']}>"
    msg['To'] = to_email
    msg['Auto-Submitted'] = 'auto-generated'
    msg.set_content(body)

    context = ssl.create_default_context()
    try:
        if settings['use_ssl']:
            with smtplib.SMTP_SSL(settings['host'], settings['port'], context=context) as server:
                if settings['user'] and settings['password']:
                    server.login(settings['user'], settings['password'])
                server.send_message(msg)
        else:
            with smtplib.SMTP(settings['host'], settings['port']) as server:
                if settings['use_tls']:
                    server.starttls(context=context)
                if settings['user'] and settings['password']:
                    server.login(settings['user'], settings['password'])
                server.send_message(msg)
    except Exception as exc:
        return False, str(exc)

    return True, None


def verify_otp(email: str, code: str, purpose: str):
    otp = session.get('otp')
    if not otp:
        return False, 'No OTP request found.'
    if otp.get('email') != email or otp.get('purpose') != purpose:
        return False, 'OTP request does not match.'
    if time.time() > otp.get('expires_at', 0):
        return False, 'OTP has expired.'
    if otp.get('code') != code:
        return False, 'Invalid OTP.'
    return True, None


def generate_device_token():
    """Generate a secure random token for device authentication."""
    return secrets.token_urlsafe(32)


def hash_device_token(token: str):
    """Hash a device token for secure storage."""
    return hashlib.sha256(token.encode()).hexdigest()


def create_device_token(email: str):
    """Create and store a device token for a user."""
    from config import get_connection, get_db_schema
    from psycopg import sql
    
    token = generate_device_token()
    token_hash = hash_device_token(token)
    expires_at = datetime.utcnow() + timedelta(days=30)
    
    schema = get_db_schema()
    try:
        with get_connection() as conn:
            with conn.cursor() as cur:
                # Create table if it doesn't exist
                cur.execute(
                    sql.SQL(
                        """CREATE TABLE IF NOT EXISTS {}.device_tokens (
                            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                            email TEXT NOT NULL,
                            token_hash TEXT NOT NULL,
                            expires_at TIMESTAMPTZ NOT NULL,
                            created_at TIMESTAMPTZ NOT NULL DEFAULT now()
                        )"""
                    ).format(sql.Identifier(schema))
                )
                
                # Store the token
                cur.execute(
                    sql.SQL(
                        """INSERT INTO {}.device_tokens (email, token_hash, expires_at)
                           VALUES (%s, %s, %s)"""
                    ).format(sql.Identifier(schema)),
                    [email.lower().strip(), token_hash, expires_at]
                )
        return token
    except Exception as e:
        print(f"Error creating device token: {e}")
        return None


def verify_device_token(token: str):
    """Verify a device token and return the associated email if valid."""
    from config import get_connection, get_db_schema
    from psycopg import sql
    
    if not token:
        return None
    
    token_hash = hash_device_token(token)
    schema = get_db_schema()
    
    try:
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    sql.SQL(
                        """SELECT email FROM {}.device_tokens
                           WHERE token_hash = %s AND expires_at > now()
                           ORDER BY created_at DESC LIMIT 1"""
                    ).format(sql.Identifier(schema)),
                    [token_hash]
                )
                row = cur.fetchone()
                if row:
                    return row[0]
    except Exception as e:
        print(f"Error verifying device token: {e}")
    
    return None


def cleanup_expired_device_tokens():
    """Remove expired device tokens from the database."""
    from config import get_connection, get_db_schema
    from psycopg import sql
    
    schema = get_db_schema()
    try:
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    sql.SQL("DELETE FROM {}.device_tokens WHERE expires_at < now()").format(sql.Identifier(schema))
                )
    except Exception:
        pass

# Configure Google AI - environment variable only
def get_google_api_key():
    # Accept either GOOGLE_API_KEY or GEMINI_API_KEY and normalize accidental quotes/spaces.
    key = os.environ.get('GOOGLE_API_KEY') or os.environ.get('GEMINI_API_KEY') or ''
    return key.strip().strip('"').strip("'")

@app.route('/')
def index():
    return redirect(url_for('login'))


@app.route('/chatbot')
def chatbot_route():
    # explicit route for /chatbot for convenience
    requested_restaurant_id = request.args.get('restaurant_id')
    if requested_restaurant_id:
        session['restaurant_id'] = requested_restaurant_id
    restaurant_id = requested_restaurant_id or get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/chatbot.html', cfg=cfg)

@app.route('/chat', methods=['POST'])
@login_required
def admin_chat():
    """Handle admin chat messages for the help bot."""
    data = request.get_json()
    user_message = data.get('message', '')
    
    if not user_message:
        return jsonify({'error': 'No message provided'}), 400
    
    # Get fresh API key in case it was just configured
    api_key = get_google_api_key()
    if not api_key:
        return jsonify({'reply': 'AI is not configured. Please add your Google API key in the settings.'}), 200
    
    client = genai.Client(api_key=api_key)
    
    try:
        # Load config to get restaurant context
        restaurant_id = get_current_restaurant_id()
        cfg = load_config(restaurant_id)
        establishment_name = cfg.get('establishment_name', 'your restaurant')
        
        # Create admin-specific prompt
        system_prompt = f"""You are a helpful AI assistant for restaurant administrators.
You help with:
- Understanding the admin dashboard features
- Menu management questions
- Settings configuration
- AI training tips
- General restaurant management advice

Restaurant name: {establishment_name}

Formatting Guidelines:
- Use **bold** for important terms and section headers
- Use `code blocks` for UI element names (like buttons, menu items)
- Use bullet points (-) for lists
- Use numbered lists (1., 2., 3.) for step-by-step instructions
- Keep responses concise and well-structured

Respond in a friendly, helpful manner. Keep responses concise and focused on helping the administrator."""
        
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                {"role": "user", "parts": [{"text": user_message}]}
            ],
            config=genai.types.GenerateContentConfig(
                system_instruction=system_prompt,
                temperature=0.7,
                max_output_tokens=500,
                top_p=0.9,
                top_k=40
            )
        )
        
        return jsonify({'reply': response.text})
    
    except Exception as e:
        return jsonify({'reply': f'Sorry, I encountered an error: {str(e)}'}), 500


@app.route('/admin', methods=['GET', 'POST'])
def admin():
    # keep legacy /admin route but redirect to the new /superadmin URL
    return redirect(url_for('superadmin'))


@app.route('/api/superadmin/auth/status', methods=['GET'])
def api_superadmin_auth_status():
    return jsonify({'authenticated': _is_superadmin_authenticated()})


@app.route('/api/superadmin/auth/login', methods=['POST'])
def api_superadmin_auth_login():
    data = request.get_json(silent=True) or {}
    username = data.get('username', '')
    password = data.get('password', '')

    if _verify_superadmin_credentials(username, password):
        _set_superadmin_authenticated(True)
        return jsonify({'ok': True, 'authenticated': True})

    _set_superadmin_authenticated(False)
    return jsonify({
        'ok': False,
        'authenticated': False,
        'message': 'Invalid superadmin credentials.'
    }), 401


@app.route('/api/superadmin/auth/logout', methods=['POST'])
def api_superadmin_auth_logout():
    _set_superadmin_authenticated(False)
    return jsonify({'ok': True, 'authenticated': False})


@app.route('/superadmin', methods=['GET', 'POST'])
def superadmin():
    # Clear restaurant_id when entering superadmin mode
    if 'restaurant_id' in session:
        session.pop('restaurant_id')

    # Force re-authentication on every page load/refresh.
    if request.method == 'GET':
        _set_superadmin_authenticated(False)
    
    if request.method == 'POST':
        config_val = request.form.get('config', '')
        save_config({'config': config_val})
        return redirect(url_for('superadmin'))
    return render_template(
        'superadmin/superadmin.html',
        superadmin_authenticated=_is_superadmin_authenticated()
    )


@app.route('/api/superadmin/restaurants', methods=['GET'])
@superadmin_required
def api_superadmin_restaurants():
    """List all registered restaurants."""
    from tools import get_all_restaurants
    restaurants = get_all_restaurants()
    return jsonify({'restaurants': restaurants})


@app.route('/api/superadmin/stats', methods=['GET'])
@superadmin_required
def api_superadmin_stats():
    """Get platform-wide statistics."""
    from tools import get_platform_stats
    stats = get_platform_stats()
    return jsonify(stats)


@app.route('/api/superadmin/tenant/<restaurant_id>/details', methods=['GET'])
@superadmin_required
def api_superadmin_tenant_details(restaurant_id):
    """Get detailed tenant data for superadmin modal view."""
    from tools import get_restaurant_details

    days_raw = request.args.get('days', '14')
    try:
        days = int(days_raw)
    except Exception:
        days = 14

    details = get_restaurant_details(restaurant_id, days=days)
    if not details:
        return jsonify({'error': 'Tenant not found'}), 404
    return jsonify(details)


@app.route('/api/superadmin/tenant/<restaurant_id>/delete', methods=['POST'])
@superadmin_required
def api_superadmin_delete_tenant(restaurant_id):
    """Delete a tenant and all tenant-scoped data after password reconfirmation."""
    from tools import delete_tenant_data

    payload = request.get_json(silent=True) or {}
    password = payload.get('password', '')
    if not _verify_superadmin_password(password):
        return jsonify({'success': False, 'message': 'Password confirmation failed.'}), 403

    result = delete_tenant_data(restaurant_id)
    if not result.get('success'):
        return jsonify({'success': False, 'message': result.get('message', 'Unable to delete tenant.')}), 500

    if session.get('restaurant_id') == restaurant_id:
        session.pop('restaurant_id', None)

    return jsonify({'success': True, 'result': result})


@app.route('/api/superadmin/manage/<restaurant_id>', methods=['POST'])
@superadmin_required
def api_superadmin_manage_restaurant(restaurant_id):
    """Switch to manage a specific restaurant."""
    # Store the current superadmin status before switching
    if not session.get('is_superadmin'):
        session['is_superadmin'] = True
    
    # Set the restaurant_id in session
    session['restaurant_id'] = restaurant_id
    return jsonify({'success': True, 'redirect': '/dashboard'})


@app.route('/api/superadmin/system-prompt', methods=['GET'])
@superadmin_required
def api_get_global_system_prompt():
    """Get the global system prompt."""
    from tools import load_global_system_prompt
    prompt = load_global_system_prompt()
    return jsonify({'prompt': prompt})


@app.route('/api/superadmin/system-prompt', methods=['POST'])
@superadmin_required
def api_save_global_system_prompt():
    """Save the global system prompt."""
    from tools import save_global_system_prompt
    data = request.get_json() or {}
    prompt = data.get('prompt', '')
    success = save_global_system_prompt(prompt)
    return jsonify({'success': success})


@app.route('/admin-client', methods=['GET', 'POST'])
@login_required
def admin_client():
    if request.method == 'POST':
        restaurant_id = get_current_restaurant_id()
        data = {
            'establishment_name': request.form.get('establishment_name', ''),
            'logo_url': request.form.get('logo_url', ''),
            'color_hex': request.form.get('color_hex', ''),
            'font_family': request.form.get('font_family', ''),
            'menu_text': request.form.get('menu_text', ''),
            'image_urls': [u.strip() for u in request.form.get('image_urls', '').splitlines() if u.strip()]
        }
        save_config(data, restaurant_id)
        return redirect(url_for('admin_client'))
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/admin-client.html', cfg=cfg)

@app.route('/dashboard')
@login_required
def dashboard():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/dashboard.html', cfg=cfg)

@app.route('/orders')
@login_required
def orders():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/orders.html', cfg=cfg)

@app.route('/kitchen')
@login_required
def kitchen():
    """Kitchen/Cashier orders tracking page."""
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/kitchen.html', cfg=cfg)

@app.route('/menu')
@login_required
def menu():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/menu.html', cfg=cfg)


@app.route('/menu/add', methods=['POST'])
@login_required
def menu_add_item():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    name = request.form.get('name', '').strip()
    description = request.form.get('description', '').strip()
    price = _strip_currency_tokens(request.form.get('price'))
    description, price = _apply_variant_form_values(description, price, request.form)
    known_categories = _extract_known_categories(cfg.get('menu_items', []))
    category_input = request.form.get('category', '').strip()
    category = category_input or infer_menu_category(name, description, known_categories)
    normalized_name = normalize_menu_item_name(name, category)
    if normalized_name:
        name = normalized_name
    status = request.form.get('status', '').strip() or 'Live'
    image_url = request.form.get('image_url', '').strip()

    if not name:
        return redirect(url_for('menu'))

    items = cfg.get('menu_items', [])
    item_id = str(uuid.uuid4())
    
    # Store the item in database with ID
    schema = get_db_schema()
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                psycopg.sql.SQL(
                    """
                    INSERT INTO {}.menu_items (id, restaurant_id, name, description, price, category, status, image_url)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    """
                ).format(psycopg.sql.Identifier(schema)),
                [item_id, restaurant_id, name, description, price, category, status, image_url]
            )
    
    items.append({
        'id': item_id,
        'name': name,
        'description': description,
        'price': price,
        'category': category,
        'status': status,
        'image_url': image_url
    })
    cfg['menu_items'] = items
    save_config(cfg, restaurant_id)
    return redirect(url_for('menu'))


@app.route('/menu/update', methods=['POST'])
@login_required
def menu_update_item():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    index_raw = request.form.get('item_index', '').strip()
    try:
        index = int(index_raw)
    except ValueError:
        return redirect(url_for('menu'))

    items = cfg.get('menu_items', [])
    if index < 0 or index >= len(items):
        return redirect(url_for('menu'))

    name = request.form.get('name', '').strip()
    if not name:
        return redirect(url_for('menu'))

    description = request.form.get('description', '').strip()
    incoming_price = _strip_currency_tokens(request.form.get('price'))
    description, final_price = _apply_variant_form_values(description, incoming_price, request.form)
    known_categories = _extract_known_categories(items)
    category_input = request.form.get('category', '').strip()
    category = category_input or infer_menu_category(name, description, known_categories)
    normalized_name = normalize_menu_item_name(name, category)
    if normalized_name:
        name = normalized_name

    # Preserve ID from existing item
    existing_item = items[index]
    item_id = existing_item.get('id') or str(uuid.uuid4())
    
    updated_item = {
        'id': item_id,
        'name': name,
        'description': description,
        'price': final_price,
        'category': category,
        'status': request.form.get('status', '').strip() or 'Live',
        'image_url': request.form.get('image_url', '').strip()
    }
    
    # Update database
    schema = get_db_schema()
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                psycopg.sql.SQL(
                    """
                    UPDATE {}.menu_items
                    SET name = %s, description = %s, price = %s, category = %s, status = %s, image_url = %s, updated_at = now()
                    WHERE id = %s AND restaurant_id = %s
                    """
                ).format(psycopg.sql.Identifier(schema)),
                [name, description, updated_item['price'], category, updated_item['status'], 
                 updated_item['image_url'], item_id, restaurant_id]
            )
    
    items[index] = updated_item
    cfg['menu_items'] = items
    save_config(cfg, restaurant_id)
    return redirect(url_for('menu'))


@app.route('/menu/bulk-category-update', methods=['POST'])
@login_required
def menu_bulk_category_update():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    
    # Get list of indices from form. Expecting multiple values for 'item_indices'
    indices_raw = request.form.getlist('item_indices')
    new_category = request.form.get('category', '').strip()
    
    if not indices_raw or not new_category:
        return redirect(url_for('menu'))

    menu_items = cfg.get('menu_items', [])
    updated_count = 0
    
    # Sort indices in reverse order if we were deleting, but for update order doesn't matter much unless we validate exists.
    # We should validate indices are integers.
    try:
        indices = [int(idx) for idx in indices_raw]
    except ValueError:
        return redirect(url_for('menu'))
        
    for index in indices:
        if 0 <= index < len(menu_items):
            menu_items[index]['category'] = new_category
            updated_count += 1
            
    if updated_count > 0:
        cfg['menu_items'] = menu_items
        save_config(cfg, restaurant_id)
        
    return redirect(url_for('menu'))


@app.route('/menu/auto-categorize-uncategorized', methods=['POST'])
@login_required
def menu_auto_categorize_uncategorized():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    menu_items = cfg.get('menu_items', [])

    uncategorized_rows = []
    for idx, item in enumerate(menu_items):
        category = (item.get('category') or '').strip().lower()
        if not category or category == 'uncategorized':
            uncategorized_rows.append((idx, item))

    if not uncategorized_rows:
        return jsonify({'updated': 0, 'total_uncategorized': 0, 'message': 'No uncategorized items found.'})

    known_categories = _extract_known_categories(menu_items)
    allowed_categories = [c for c in known_categories if c.strip().lower() != 'uncategorized']
    for default_cat in ['Appetizers', 'Main Course', 'Desserts', 'Drinks/Beverages']:
        if default_cat not in allowed_categories:
            allowed_categories.append(default_cat)

    payload_items = []
    for idx, item in uncategorized_rows:
        payload_items.append({
            'index': idx,
            'name': (item.get('name') or '').strip(),
            'description': (item.get('description') or '').strip(),
            'price': (item.get('price') or '').strip()
        })

    def fallback_categorize(reason: str):
        updated_local = 0
        for idx, item in uncategorized_rows:
            current = (menu_items[idx].get('category') or '').strip().lower()
            if current and current != 'uncategorized':
                continue
            normalized = infer_menu_category(item.get('name', ''), item.get('description', ''), allowed_categories)
            if not normalized or normalized.strip().lower() == 'uncategorized':
                normalized = 'Main Course'
            menu_items[idx]['category'] = normalized
            updated_local += 1

        if updated_local > 0:
            cfg['menu_items'] = menu_items
            save_config(cfg, restaurant_id)

        return jsonify({
            'updated': updated_local,
            'total_uncategorized': len(uncategorized_rows),
            'categories': allowed_categories,
            'fallback': True,
            'warning': reason
        })

    api_key = get_google_api_key()
    if not api_key:
        return fallback_categorize('Google API key not configured. Used local categorization instead.')

    try:
        client = genai.Client(api_key=api_key)
        system_instruction = (
            'You categorize restaurant menu items. '
            'Return JSON only: an array of objects with keys index and category. '
            'Use only categories from the provided allowed_categories list. '
            'Do not invent categories. '
            'If unsure, choose the closest allowed category.'
        )
        user_payload = {
            'allowed_categories': allowed_categories,
            'items': payload_items
        }

        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[{"role": "user", "parts": [{"text": json.dumps(user_payload)}]}],
            config=genai.types.GenerateContentConfig(
                system_instruction=system_instruction,
                response_mime_type='application/json',
                temperature=0.1,
                max_output_tokens=6000
            )
        )

        text = (response.text or '').strip()
        if not text:
            raise ValueError('Empty AI response')
        if not text.startswith('['):
            start = text.find('[')
            end = text.rfind(']')
            if start != -1 and end != -1 and end > start:
                text = text[start:end + 1]

        suggestions = json.loads(text)
        if not isinstance(suggestions, list):
            raise ValueError('AI response is not a list')

        allowed_lookup = {c.lower(): c for c in allowed_categories}
        classified_indices = set()
        updated = 0
        for row in suggestions:
            if not isinstance(row, dict):
                continue
            idx = row.get('index')
            category_raw = (row.get('category') or '').strip()
            if not isinstance(idx, int):
                continue
            if idx < 0 or idx >= len(menu_items):
                continue
            classified_indices.add(idx)

            normalized = allowed_lookup.get(category_raw.lower())
            if not normalized:
                item = menu_items[idx]
                normalized = infer_menu_category(item.get('name', ''), item.get('description', ''), allowed_categories)
                if normalized.lower() == 'uncategorized':
                    normalized = 'Main Course'

            current = (menu_items[idx].get('category') or '').strip().lower()
            if not current or current == 'uncategorized':
                menu_items[idx]['category'] = normalized
                updated += 1

        # Fill any missing AI classifications with local categorization.
        for idx, item in uncategorized_rows:
            if idx in classified_indices:
                continue
            current = (menu_items[idx].get('category') or '').strip().lower()
            if current and current != 'uncategorized':
                continue
            normalized = infer_menu_category(item.get('name', ''), item.get('description', ''), allowed_categories)
            if not normalized or normalized.strip().lower() == 'uncategorized':
                normalized = 'Main Course'
            menu_items[idx]['category'] = normalized
            updated += 1

        if updated > 0:
            cfg['menu_items'] = menu_items
            save_config(cfg, restaurant_id)

        return jsonify({
            'updated': updated,
            'total_uncategorized': len(uncategorized_rows),
            'categories': allowed_categories
        })
    except Exception as exc:
        app.logger.exception('Auto-categorization failed; using fallback')
        err = str(exc)
        if 'API_KEY_INVALID' in err or 'API key not valid' in err or 'INVALID_ARGUMENT' in err:
            reason = 'Gemini API key is invalid. Used local categorization instead.'
        else:
            reason = 'Gemini is unavailable right now. Used local categorization instead.'
        return fallback_categorize(reason)


def _normalize_menu_key(value: str) -> str:
    text = (value or '').strip().lower()
    return re.sub(r'[^a-z0-9]+', '', text)


@app.route('/menu/upload', methods=['POST'])
@login_required
def menu_upload_menu_file():
    try:
        restaurant_id = get_current_restaurant_id()
        file = request.files.get('menu_file')
        merge_mode = request.form.get('merge_mode', 'replace') == 'true'
        
        if not file or not file.filename:
            return jsonify({'error': 'No file provided'}), 400
        filename = file.filename.lower()
        is_txt = filename.endswith('.txt')
        is_pdf = filename.endswith('.pdf')
        is_png = filename.endswith('.png')
        if not (is_txt or is_pdf or is_png):
            return jsonify({'error': 'Only .txt, .pdf, or .png files are supported'}), 400

        data_bytes = file.read()
        content = ''
        
        if is_pdf:
            content = extract_pdf_text(data_bytes)
            if not content:
                return jsonify({'error': 'Unable to extract text from PDF'}), 400
        elif is_png:
            mime_type = file.content_type or 'image/png'
            content = extract_image_text_with_ai(data_bytes, mime_type)
            if not content:
                return jsonify({'error': 'Unable to extract text from PNG. The AI service may be unavailable or the image may not contain readable text. Please try again or use a .txt or .pdf file instead.'}), 400
        else:
            content = data_bytes.decode('utf-8', errors='ignore')

        cfg = load_config(restaurant_id)

        # Parse the extracted text into structured items
        items = parse_menu_txt_with_ai(content)
        if not items:
            items = parse_menu_txt(content)
        
        if not items:
            return jsonify({'error': 'No menu items found in file'}), 400

        # Deduplicate items by name (case-insensitive)
        seen_names = set()
        unique_items = []
        for item in items:
            name_lower = (item.get('name') or '').strip().lower()
            if name_lower and name_lower not in seen_names:
                seen_names.add(name_lower)
                unique_items.append(item)

        known_categories = _extract_known_categories(cfg.get('menu_items', []))
        existing_items = cfg.get('menu_items', [])
        image_map = {}
        for item in existing_items:
            key = _normalize_menu_key(item.get('name', ''))
            if key and item.get('image_url'):
                image_map[key] = item.get('image_url')

        normalized_items = []
        review_flagged = 0
        for item in unique_items:
            normalized_item, needs_review = _normalize_imported_menu_item(item, known_categories)
            if not normalized_item:
                continue

            category = (normalized_item.get('category') or '').strip()
            if category and category not in known_categories and category.lower() != 'uncategorized':
                known_categories.append(category)

            key = _normalize_menu_key(normalized_item.get('name', ''))
            if key and key in image_map and not normalized_item.get('image_url'):
                normalized_item['image_url'] = image_map[key]

            if needs_review:
                review_flagged += 1
            normalized_items.append(normalized_item)

        if not normalized_items:
            return jsonify({'error': 'No valid menu items found after normalization'}), 400

        # Handle merge mode
        if merge_mode:
            # Merge with existing items - detect similar items and update, add new ones
            merged_items = list(existing_items)
            existing_names_normalized = {_normalize_menu_key(item.get('name', '')): i for i, item in enumerate(merged_items)}
            
            added_count = 0
            updated_count = 0
            
            for new_item in normalized_items:
                new_key = _normalize_menu_key(new_item.get('name', ''))
                if new_key in existing_names_normalized:
                    # Item already exists - update it with new info
                    idx = existing_names_normalized[new_key]
                    # Preserve ID and image if exists
                    if 'id' in merged_items[idx]:
                        new_item['id'] = merged_items[idx]['id']
                    if not new_item.get('image_url') and merged_items[idx].get('image_url'):
                        new_item['image_url'] = merged_items[idx].get('image_url')
                    if not new_item.get('image_data') and merged_items[idx].get('image_data'):
                        new_item['image_data'] = merged_items[idx].get('image_data')
                    merged_items[idx] = new_item
                    updated_count += 1
                else:
                    # New item - add it
                    merged_items.append(new_item)
                    added_count += 1
            
            cfg['menu_items'] = merged_items
            result_message = f"Added {added_count} new items, updated {updated_count} existing items"
        else:
            # Replace mode (default)
            cfg['menu_items'] = normalized_items
            result_message = f"Replaced menu with {len(normalized_items)} items"
        
        save_config(cfg, restaurant_id)
        save_training_upload_bytes(restaurant_id, file.filename, data_bytes)

        return jsonify({
            'saved': len(normalized_items),
            'flagged_for_review': review_flagged,
            'message': result_message,
            'merge_mode': merge_mode,
        })
    except Exception as exc:
        app.logger.exception('Menu upload failed')
        return jsonify({'error': 'Menu upload failed', 'detail': str(exc)}), 500


@app.route('/menu/photos/upload', methods=['POST'])
@login_required
def menu_upload_photos():
    try:
        restaurant_id = get_current_restaurant_id()
        files = request.files.getlist('menu_photos')
        if not files:
            return jsonify({'error': 'No files provided'}), 400

        cfg = load_config(restaurant_id)
        items = cfg.get('menu_items', [])
        name_index = {}
        for idx, item in enumerate(items):
            key = _normalize_menu_key(item.get('name', ''))
            if key:
                name_index[key] = idx

        matched = 0
        unmatched = []
        uploaded_count = 0

        schema = get_db_schema()
        for file in files:
            if not file or not file.filename:
                continue
            if not allowed_file(file.filename):
                unmatched.append(file.filename)
                continue

            safe_name = secure_filename(file.filename)
            if not safe_name:
                unmatched.append(file.filename)
                continue

            # Read file bytes and get MIME type
            mime_type = file.content_type or 'application/octet-stream'
            file_bytes = file.read()
            uploaded_count += 1

            # Match filename to menu item
            key = _normalize_menu_key(Path(file.filename).stem)
            if key in name_index:
                idx = name_index[key]
                item_id = items[idx].get('id')
                
                if item_id:
                    # Store in database
                    with get_connection() as conn:
                        with conn.cursor() as cur:
                            cur.execute(
                                psycopg.sql.SQL(
                                    """
                                    UPDATE {}.menu_items
                                    SET image_data = %s, image_mime = %s
                                    WHERE id = %s AND restaurant_id = %s
                                    """
                                ).format(psycopg.sql.Identifier(schema)),
                                [file_bytes, mime_type, item_id, restaurant_id]
                            )
                matched += 1
            else:
                unmatched.append(file.filename)

        cfg['menu_items'] = items
        save_config(cfg, restaurant_id)

        return jsonify({
            'uploaded': uploaded_count,
            'matched': matched,
            'unmatched': unmatched
        })
    except Exception as exc:
        app.logger.exception('Menu photo upload failed')
        return jsonify({'error': 'Menu photo upload failed', 'detail': str(exc)}), 500


# Menu Photo Serving Route
@app.route('/menu/photo/<photo_id>', methods=['GET'])
@login_required
def menu_get_photo(photo_id):
    """Serve a menu item photo from the database."""
    try:
        restaurant_id = get_current_restaurant_id()
        schema = get_db_schema()
        
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    psycopg.sql.SQL(
                        """
                        SELECT image_data, image_mime
                        FROM {}.menu_items
                        WHERE id = %s AND restaurant_id = %s
                        """
                    ).format(psycopg.sql.Identifier(schema)),
                    [photo_id, restaurant_id]
                )
                row = cur.fetchone()

        if not row or not row[0]:
            return jsonify({'error': 'Photo not found'}), 404

        image_data = row[0]
        mime_type = row[1] or 'image/jpeg'

        response = Response(image_data, mimetype=mime_type)
        response.headers['Cache-Control'] = 'public, max-age=31536000'
        return response
    except Exception as exc:
        app.logger.exception('Failed to fetch menu photo')
        return jsonify({'error': 'Failed to fetch photo'}), 500


@app.route('/brand/image/<image_kind>/<restaurant_id>', methods=['GET'])
def brand_get_image(image_kind, restaurant_id):
    """Serve a brand logo/avatar image from the database."""
    try:
        if image_kind not in BRAND_IMAGE_COLUMNS or not restaurant_id:
            return jsonify({'error': 'Image not found'}), 404

        data_col, mime_col = BRAND_IMAGE_COLUMNS[image_kind]
        schema = get_db_schema()

        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    psycopg.sql.SQL(
                        "SELECT {}, {} FROM {}.brand_settings WHERE restaurant_id = %s"
                    ).format(
                        psycopg.sql.Identifier(data_col),
                        psycopg.sql.Identifier(mime_col),
                        psycopg.sql.Identifier(schema)
                    ),
                    [restaurant_id]
                )
                row = cur.fetchone()

        if not row or not row[0]:
            return jsonify({'error': 'Image not found'}), 404

        image_data = row[0]
        if isinstance(image_data, memoryview):
            image_data = image_data.tobytes()

        mime_type = _normalize_image_mime(row[1] if len(row) > 1 else '', image_data)
        response = Response(image_data, mimetype=mime_type)
        response.headers['Cache-Control'] = 'no-store, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response
    except Exception:
        app.logger.exception('Failed to fetch brand image')
        return jsonify({'error': 'Failed to fetch image'}), 500

@app.route('/customers')
@login_required
def customers():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/customers.html', cfg=cfg)

@app.route('/reports')
@login_required
def reports():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/reports.html', cfg=cfg)

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)

    def _normalize_hex_color(value, fallback=''):
        raw = (value or '').strip().lower()
        if not raw:
            return fallback
        if not raw.startswith('#'):
            raw = '#' + raw
        if re.match(r'^#[0-9a-f]{3}$', raw):
            raw = '#' + ''.join(ch * 2 for ch in raw[1:])
        if re.match(r'^#[0-9a-f]{8}$', raw):
            raw = '#' + raw[1:7]
        if re.match(r'^#[0-9a-f]{6}$', raw):
            return raw
        return fallback

    def _auto_foreground(bg_color: str, light: str = '#ffffff', dark: str = '#000000') -> str:
        color = _normalize_hex_color(bg_color, '#1e40af')
        hex_value = color[1:]
        r = int(hex_value[0:2], 16) / 255.0
        g = int(hex_value[2:4], 16) / 255.0
        b = int(hex_value[4:6], 16) / 255.0

        def _lin(channel):
            return channel / 12.92 if channel <= 0.03928 else ((channel + 0.055) / 1.055) ** 2.4

        luminance = 0.2126 * _lin(r) + 0.7152 * _lin(g) + 0.0722 * _lin(b)
        return dark if luminance > 0.5 else light

    if request.method == 'POST':
        # collect branding & display fields and merge into existing config
        currency_choice = request.form.get('currency_choice', '').strip()
        currency_code = cfg.get('currency_code', 'PHP')
        currency_symbol = cfg.get('currency_symbol', '₱')
        currency_map = {
            'PHP': '₱',
            'USD': '$',
            'EUR': '€',
            'GBP': '£',
            'JPY': '¥',
            'AUD': 'A$',
            'CAD': 'C$',
            'SGD': 'S$',
            'INR': '₹'
        }
        if currency_choice:
            if '|' in currency_choice:
                parts = currency_choice.split('|', 1)
                currency_code = (parts[0] or '').strip() or currency_code
                currency_symbol = (parts[1] or '').strip() or currency_symbol
            else:
                currency_code = currency_choice
                currency_symbol = currency_map.get(currency_choice, currency_symbol)

        main_color = _normalize_hex_color(
            request.form.get('main_color', request.form.get('color_hex', cfg.get('main_color', cfg.get('color_hex', '')))),
            cfg.get('main_color', cfg.get('color_hex', '#1e40af'))
        )
        sub_color = _normalize_hex_color(
            request.form.get('sub_color', cfg.get('sub_color', '')),
            cfg.get('sub_color', '#ffd41d')
        )

        main_foreground = _auto_foreground(main_color)
        sub_foreground = _auto_foreground(sub_color)

        data = {
            'establishment_name': request.form.get('establishment_name', cfg.get('establishment_name', '')),
            'logo_url': request.form.get('logo_url', cfg.get('logo_url', '')),
            'main_color': main_color,
            'main_foreground': main_foreground,
            'sub_color': sub_color,
            'sub_foreground': sub_foreground,
            'font_family': request.form.get('font_family', cfg.get('font_family', '')),
            'menu_text': request.form.get('menu_text', cfg.get('menu_text', '')),
            'image_urls': [u.strip() for u in request.form.get('image_urls', "\n".join(cfg.get('image_urls', []))).splitlines() if u.strip()],
            'open_time': request.form.get('open_time', cfg.get('open_time', '')),
            'close_time': request.form.get('close_time', cfg.get('close_time', '')),
            'tax_rate': request.form.get('tax_rate', cfg.get('tax_rate', '')),
            'currency_code': currency_code,
            'currency_symbol': currency_symbol
        }
        # merge with existing config to avoid wiping other keys
        # handle logo upload (optional)
        logo_url = request.form.get('logo_url', cfg.get('logo_url', ''))
        logo_file = request.files.get('logo_file')
        if logo_file and logo_file.filename:
            if allowed_file(logo_file.filename):
                try:
                    logo_url = _store_brand_image_upload(restaurant_id, 'logo', logo_file)
                    data['logo_uploaded_by'] = session.get('user')
                    data['logo_uploaded_at'] = datetime.now(timezone.utc)
                except Exception:
                    app.logger.exception('Logo upload failed')
                    flash('Failed to upload logo.')
            else:
                flash('Unsupported logo file type.')
        else:
            try:
                logo_url = _migrate_static_brand_image(restaurant_id, 'logo', logo_url)
            except Exception:
                app.logger.exception('Failed migrating legacy logo image')

        data['logo_url'] = logo_url

        # handle chatbot avatar upload (optional)
        avatar_url = request.form.get('chatbot_avatar_url', cfg.get('chatbot_avatar', ''))
        avatar_file = request.files.get('chatbot_avatar_file')
        if avatar_file and avatar_file.filename:
            if allowed_file(avatar_file.filename):
                try:
                    avatar_url = _store_brand_image_upload(restaurant_id, 'chatbot_avatar', avatar_file)
                    data['chatbot_avatar_uploaded_by'] = session.get('user')
                    data['chatbot_avatar_uploaded_at'] = datetime.now(timezone.utc)
                except Exception:
                    app.logger.exception('Chatbot avatar upload failed')
                    flash('Failed to upload chatbot avatar.')
            else:
                flash('Unsupported avatar file type.')
        else:
            try:
                avatar_url = _migrate_static_brand_image(restaurant_id, 'chatbot_avatar', avatar_url)
            except Exception:
                app.logger.exception('Failed migrating legacy chatbot avatar')

        data['chatbot_avatar'] = avatar_url

        # Persist only the POST-normalized payload. Avoid re-saving stale blob fields
        # (logo_data/chatbot_avatar_data) from the pre-request cfg snapshot.
        save_config(data, restaurant_id)
        return redirect(url_for('settings'))
    return render_template('clients/settings.html', cfg=cfg)


@app.route('/settings/clear-menu', methods=['POST'])
@login_required
def settings_clear_menu():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    cfg['menu_items'] = []
    save_config(cfg, restaurant_id)
    return jsonify({'cleared': True})

@app.route('/ai-training', methods=['GET', 'POST'])
@login_required
def ai_training():
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/ai-training.html', cfg=cfg)


@app.route('/ai-training/files', methods=['GET'])
@login_required
def ai_training_files():
    restaurant_id = get_current_restaurant_id()
    entries = load_training_manifest(restaurant_id)
    training_dir = get_training_dir(restaurant_id)

    filtered = []
    for entry in entries:
        stored_name = entry.get('stored_name')
        if not stored_name:
            continue
        file_path = training_dir / stored_name
        if not file_path.exists():
            continue
        size_bytes = file_path.stat().st_size
        filtered.append({
            'id': entry.get('id'),
            'original_name': entry.get('original_name'),
            'stored_name': stored_name,
            'size_bytes': size_bytes,
            'uploaded_at': entry.get('uploaded_at'),
            'status': entry.get('status', 'ready')
        })

    filtered.sort(key=lambda x: x.get('uploaded_at') or '', reverse=True)
    return jsonify({'files': filtered})


@app.route('/ai-training/upload', methods=['POST'])
@login_required
def ai_training_upload():
    restaurant_id = get_current_restaurant_id()
    if 'training_files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('training_files')
    if not files:
        return jsonify({'error': 'No files provided'}), 400

    entries = load_training_manifest(restaurant_id)
    cfg = load_config(restaurant_id)
    configured_currency_code = cfg.get('currency_code', 'PHP')
    configured_currency_symbol = cfg.get('currency_symbol', '₱')
    training_dir = get_training_dir(restaurant_id)
    saved = []
    errors = []
    currency_warnings = []

    for file in files:
        filename = file.filename or ''
        if not filename:
            continue
        started_at = datetime.now(timezone.utc).isoformat()
        start_perf = time.perf_counter()
        if not training_allowed_file(filename):
            errors.append({'file': filename, 'error': 'Unsupported file type'})
            add_training_history_entry(restaurant_id, {
                'id': uuid.uuid4().hex,
                'action': f'File upload: {filename}',
                'status': 'failed',
                'started_at': started_at,
                'ended_at': datetime.now(timezone.utc).isoformat(),
                'duration_ms': int((time.perf_counter() - start_perf) * 1000)
            })
            continue
        file.seek(0, os.SEEK_END)
        size_bytes = file.tell()
        file.seek(0)
        if size_bytes > MAX_TRAINING_FILE_MB * 1024 * 1024:
            errors.append({'file': filename, 'error': 'File too large'})
            add_training_history_entry(restaurant_id, {
                'id': uuid.uuid4().hex,
                'action': f'File upload: {filename}',
                'status': 'failed',
                'started_at': started_at,
                'ended_at': datetime.now(timezone.utc).isoformat(),
                'duration_ms': int((time.perf_counter() - start_perf) * 1000)
            })
            continue

        safe_name = secure_filename(filename)
        ext = Path(safe_name).suffix.lower()
        stored_name = f"{uuid.uuid4().hex}{ext}"
        dest = training_dir / stored_name
        file.save(str(dest))

        entry = {
            'id': uuid.uuid4().hex,
            'original_name': safe_name,
            'stored_name': stored_name,
            'uploaded_at': datetime.now(timezone.utc).isoformat(),
            'status': 'ready',
            'size_bytes': int(size_bytes)
        }

        # Unified two-pass ingestion for training files:
        # pass 1 = extract plain text from file, pass 2 = AI structure/tagging.
        try:
            preview_text = _build_training_preview_text(dest)
            if ext == '.pdf' and preview_text:
                detected_currency = _detect_currency_from_text(preview_text)
                mismatch_warning = _build_currency_mismatch_warning(
                    configured_currency_code,
                    configured_currency_symbol,
                    detected_currency,
                )
                if mismatch_warning:
                    currency_warnings.append({
                        'file': safe_name,
                        'message': mismatch_warning,
                        'detected_currency': detected_currency,
                    })
            if preview_text:
                ai_profile = parse_training_text_with_ai(preview_text, safe_name)
                if ai_profile:
                    entry['ai_profile'] = ai_profile
                    if ai_profile.get('categories'):
                        entry['ai_categories'] = ai_profile.get('categories')
                    if ai_profile.get('document_type'):
                        entry['ai_document_type'] = ai_profile.get('document_type')
        except Exception:
            app.logger.exception('Training AI profiling failed for %s', safe_name)

        entries.append(entry)
        saved.append(entry)
        add_training_history_entry(restaurant_id, {
            'id': uuid.uuid4().hex,
            'action': f'File upload: {safe_name}',
            'status': 'completed',
            'started_at': started_at,
            'ended_at': datetime.now(timezone.utc).isoformat(),
            'duration_ms': int((time.perf_counter() - start_perf) * 1000)
        })

    save_training_manifest(restaurant_id, entries)
    return jsonify({'saved': saved, 'errors': errors, 'currency_warnings': currency_warnings})


@app.route('/ai-training/files/<file_id>', methods=['DELETE'])
@login_required
def ai_training_delete(file_id):
    restaurant_id = get_current_restaurant_id()
    entries = load_training_manifest(restaurant_id)
    training_dir = get_training_dir(restaurant_id)

    remaining = []
    deleted = False
    deleted_name = ''
    for entry in entries:
        if entry.get('id') != file_id:
            remaining.append(entry)
            continue
        deleted_name = entry.get('original_name') or entry.get('stored_name') or ''
        stored_name = entry.get('stored_name')
        if stored_name:
            file_path = training_dir / stored_name
            if file_path.exists():
                try:
                    file_path.unlink()
                except Exception:
                    pass
        deleted = True

    save_training_manifest(restaurant_id, remaining)
    if not deleted:
        return jsonify({'error': 'File not found'}), 404
    if deleted_name:
        add_training_history_entry(restaurant_id, {
            'id': uuid.uuid4().hex,
            'action': f'File delete: {deleted_name}',
            'status': 'completed',
            'started_at': datetime.now(timezone.utc).isoformat(),
            'ended_at': datetime.now(timezone.utc).isoformat(),
            'duration_ms': 0
        })
    return jsonify({'deleted': True})


@app.route('/ai-training/history', methods=['GET'])
@login_required
def ai_training_history():
    restaurant_id = get_current_restaurant_id()
    entries = load_training_history(restaurant_id)
    entries.sort(key=lambda item: item.get('started_at') or '', reverse=True)
    return jsonify({'history': entries[:50]})


@app.route('/ai-training/knowledge', methods=['GET'])
@login_required
def ai_training_knowledge():
    restaurant_id = get_current_restaurant_id()
    entries = load_training_manifest(restaurant_id)
    training_dir = get_training_dir(restaurant_id)

    file_summaries = []
    chunk_samples = []
    total_chunks = 0
    structured_chunks = 0
    max_samples = 60

    for entry in entries:
        stored_name = entry.get('stored_name')
        if not stored_name:
            continue

        file_path = training_dir / stored_name
        if not file_path.exists():
            continue

        chunks = build_training_chunks(restaurant_id, file_path, entry)
        if not chunks:
            continue

        source_name = entry.get('original_name') or stored_name
        file_chunk_count = len(chunks)
        total_chunks += file_chunk_count

        has_structured_meta = False
        for chunk in chunks:
            meta = chunk.get('metadata') or {}
            if meta.get('page') is not None or meta.get('section_title'):
                structured_chunks += 1
                has_structured_meta = True

            if len(chunk_samples) < max_samples:
                content = (chunk.get('content') or '').strip()
                chunk_samples.append({
                    'source_file': source_name,
                    'content_preview': content[:320],
                    'metadata': {
                        'file_ext': meta.get('file_ext'),
                        'page': meta.get('page'),
                        'section_title': meta.get('section_title'),
                        'identifier': meta.get('identifier'),
                    }
                })

        file_summaries.append({
            'source_file': source_name,
            'chunk_count': file_chunk_count,
            'chunk_mode': 'structured' if has_structured_meta else 'sliding',
            'file_ext': file_path.suffix.lower(),
        })

    return jsonify({
        'summary': {
            'file_count': len(file_summaries),
            'total_chunks': total_chunks,
            'structured_chunks': structured_chunks,
            'sliding_chunks': max(0, total_chunks - structured_chunks),
        },
        'files': file_summaries,
        'known_chunks': chunk_samples,
    })


@app.route('/ai-training/retrain', methods=['POST'])
@login_required
def ai_training_retrain():
    restaurant_id = get_current_restaurant_id()
    started_at = datetime.now(timezone.utc).isoformat()
    entry = {
        'id': uuid.uuid4().hex,
        'action': 'Retrain all models',
        'status': 'completed',
        'started_at': started_at,
        'ended_at': datetime.now(timezone.utc).isoformat(),
        'duration_ms': 0
    }
    add_training_history_entry(restaurant_id, entry)
    return jsonify({'status': 'completed', 'entry': entry})


def _build_training_preview_text(file_path: Path):
    suffix = file_path.suffix.lower()
    if suffix == '.pdf':
        return extract_pdf_text(file_path.read_bytes())
    if suffix == '.docx':
        return extract_docx_text(file_path.read_bytes())
    try:
        raw_text = file_path.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return ''
    if suffix == '.json':
        try:
            parsed = json.loads(raw_text)
            return json.dumps(parsed, indent=2)
        except Exception:
            return raw_text
    return raw_text


@app.route('/ai-training/files/<file_id>/preview', methods=['GET'])
@login_required
def ai_training_preview(file_id):
    restaurant_id = get_current_restaurant_id()
    entries = load_training_manifest(restaurant_id)
    training_dir = get_training_dir(restaurant_id)

    entry = next((e for e in entries if e.get('id') == file_id), None)
    if not entry:
        return jsonify({'error': 'File not found'}), 404

    stored_name = entry.get('stored_name')
    if not stored_name:
        return jsonify({'error': 'File not found'}), 404

    file_path = training_dir / stored_name
    if not file_path.exists():
        return jsonify({'error': 'File not found'}), 404

    preview_text = _build_training_preview_text(file_path)
    if not preview_text:
        return jsonify({'error': 'Unable to extract preview text'}), 400

    max_chars = 4000
    truncated = len(preview_text) > max_chars
    preview = preview_text[:max_chars]

    return jsonify({
        'id': entry.get('id'),
        'name': entry.get('original_name'),
        'ext': file_path.suffix.lower(),
        'preview': preview,
        'truncated': truncated
    })


@app.route('/login', methods=['GET', 'POST'])
def login():
    cfg = load_config()
    error = None
    
    # Check for device token before showing login form
    device_token = request.cookies.get('device_token')
    if device_token:
        email = verify_device_token(device_token)
        if email and user_exists(email):
            session['user'] = email
            session.pop('restaurant_id', None)
            get_current_restaurant_id()
            return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        remember_device = request.form.get('remember_device') == 'on'
        
        if not email:
            error = 'Email is required.'
        elif not user_exists(email):
            error = 'No account found with that email.'
        else:
            # Store remember_device preference in session
            session['remember_device'] = remember_device
            
            otp_code = set_otp(email, 'login')
            sent, send_error = send_otp_email(email, otp_code, 'login', cfg)
            session['otp_notice'] = 'We sent a verification code to your email.' if sent else None
            session['otp_warning'] = None if sent else f"Email delivery failed: {send_error}. Use the code shown below."
            return redirect(url_for('otp_verify_page', email=email, purpose='login'))
    # Return 422 for validation errors (Turbo requirement)
    if error:
        return render_template('auth/login.html', error=error), 422
    return render_template('auth/login.html', error=error)


@app.route('/otp-verify', methods=['GET'])
def otp_verify_page():
    """Display OTP verification page."""
    email = request.args.get('email', '').strip().lower()
    purpose = request.args.get('purpose', 'login')
    error = request.args.get('error', '')
    
    if not email or not session.get('otp'):
        return redirect(url_for('login'))
    
    otp_code = session.get('otp', {}).get('code')
    otp_hint = otp_code if should_show_otp_hint() else None
    notice = session.pop('otp_notice', None)
    warning = session.pop('otp_warning', None)
    
    return render_template(
        'auth/otp_verify.html',
        email=email,
        purpose=purpose,
        error=error,
        otp_hint=otp_hint,
        notice=notice,
        warning=warning
    )


@app.route('/login/verify', methods=['POST'])
def login_verify():
    cfg = load_config()
    email = request.form.get('email', '').strip().lower()
    code = request.form.get('otp', '').strip()
    ok, error = verify_otp(email, code, 'login')
    if ok:
        session.pop('otp', None)
        session['user'] = email
        session.pop('restaurant_id', None)
        get_current_restaurant_id()
        
        # Check if user wants to remember this device
        # Can be set either at login or at OTP verification
        remember_device = session.pop('remember_device', False) or request.form.get('remember_device') == 'on'
        
        response = make_response(redirect(url_for('dashboard')))
        
        if remember_device:
            # Create and set device token cookie
            token = create_device_token(email)
            if token:
                response.set_cookie(
                    'device_token',
                    token,
                    max_age=30*24*60*60,  # 30 days in seconds
                    httponly=True,
                    secure=request.is_secure,
                    samesite='Lax'
                )
        
        # Clean up expired tokens periodically
        cleanup_expired_device_tokens()
        
        return response
    
    # Redirect back to OTP verify page with error
    return redirect(url_for('otp_verify_page', email=email, purpose='login', error=error))


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    cfg = load_config()
    error = None
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        establishment_name = request.form.get('establishment_name', '')
        main_color = request.form.get('main_color', '')
        sub_color = request.form.get('sub_color', '')
        currency_choice = request.form.get('currency_choice', '').strip()

        currency_code = 'PHP'
        currency_symbol = '₱'
        currency_map = {
            'PHP': '₱',
            'USD': '$',
            'EUR': '€',
            'GBP': '£',
            'JPY': '¥',
            'AUD': 'A$',
            'CAD': 'C$',
            'SGD': 'S$',
            'INR': '₹'
        }
        if currency_choice:
            if '|' in currency_choice:
                code_part, symbol_part = currency_choice.split('|', 1)
                currency_code = (code_part or '').strip() or currency_code
                currency_symbol = (symbol_part or '').strip() or currency_symbol
            else:
                currency_code = currency_choice
                currency_symbol = currency_map.get(currency_choice, currency_symbol)

        existing_pending = session.get('pending_signup') or {}
        pending_restaurant_id = existing_pending.get('restaurant_id') or str(uuid.uuid4())

        # handle logo file upload
        logo_url = request.form.get('logo_url', '')
        logo_file = request.files.get('logo_file')
        if logo_file and logo_file.filename:
            if allowed_file(logo_file.filename):
                try:
                    logo_url = _store_brand_image_upload(pending_restaurant_id, 'logo', logo_file)
                except Exception:
                    app.logger.exception('Signup logo upload failed')
                    error = 'Failed to upload logo file.'
            else:
                error = 'Unsupported logo file type.'

        if not error:
            if not email:
                error = 'Email is required.'
            elif user_exists(email):
                error = 'A user with that email already exists.'
            else:
                session['pending_signup'] = {
                    'email': email,
                    'restaurant_id': pending_restaurant_id,
                    'establishment_name': establishment_name,
                    'logo_url': logo_url,
                    'main_color': main_color,
                    'sub_color': sub_color,
                    'currency_code': currency_code,
                    'currency_symbol': currency_symbol
                }
                otp_code = set_otp(email, 'signup')
                sent, send_error = send_otp_email(email, otp_code, 'signup', cfg)
                session['otp_notice'] = 'We sent a verification code to your email.' if sent else None
                session['otp_warning'] = None if sent else f"Email delivery failed: {send_error}. Use the code shown below."
                return redirect(url_for('otp_verify_page', email=email, purpose='signup'))

    # Return 422 for validation errors (Turbo requirement)
    if error:
        return render_template('auth/signup.html', error=error), 422
    return render_template('auth/signup.html', error=error)


@app.route('/signup/verify', methods=['POST'])
def signup_verify():
    cfg = load_config()
    pending = session.get('pending_signup') or {}
    email = request.form.get('email', '').strip().lower()
    if not pending or pending.get('email') != email:
        error = 'Signup session not found. Please start again.'
        # Return 422 for validation errors (Turbo requirement)
        return render_template('auth/signup.html', error=error), 422

    code = request.form.get('otp', '').strip()
    ok, error = verify_otp(email, code, 'signup')
    if not ok:
        # Redirect back to OTP verify page with error
        return redirect(url_for('otp_verify_page', email=email, purpose='signup', error=error))

    restaurant_id = pending.get('restaurant_id') or str(uuid.uuid4())
    meta = {
        'establishment_name': pending.get('establishment_name', ''),
        'logo_url': pending.get('logo_url', ''),
        'main_color': pending.get('main_color', ''),
        'sub_color': pending.get('sub_color', ''),
        'currency_code': pending.get('currency_code', 'PHP'),
        'currency_symbol': pending.get('currency_symbol', '₱'),
        'restaurant_id': restaurant_id
    }
    success = add_user(email, password=None, meta=meta)
    if not success:
        error = 'A user with that email already exists.'
        # Return 422 for validation errors (Turbo requirement)
        return render_template('auth/signup.html', error=error), 422

    cfg.update({
        'establishment_name': pending.get('establishment_name', ''),
        'logo_url': pending.get('logo_url', ''),
        'main_color': pending.get('main_color', ''),
        'sub_color': pending.get('sub_color', ''),
        'currency_code': pending.get('currency_code', 'PHP'),
        'currency_symbol': pending.get('currency_symbol', '₱')
    })
    save_config(cfg, restaurant_id)
    session.pop('pending_signup', None)
    session.pop('otp', None)
    return redirect(url_for('login'))


@app.route('/logout')
def logout():
    session.pop('user', None)
    session.pop('restaurant_id', None)
    session.pop('superadmin_authenticated', None)
    session.pop('is_superadmin', None)
    response = make_response(redirect(url_for('login')))
    # Clear device token cookie on logout
    response.set_cookie('device_token', '', expires=0)
    return response


@app.route('/qr-codes')
@login_required
def qr_codes():
    """Display the QR codes management page."""
    restaurant_id = get_current_restaurant_id()
    cfg = load_config(restaurant_id)
    return render_template('clients/qr-codes.html', cfg=cfg)


def _detect_lan_ip() -> str:
    """Best-effort detection of the server's LAN IPv4 address."""
    sock = None
    try:
        sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        # No packets are sent; this is used to select the active outbound interface.
        sock.connect(('8.8.8.8', 80))
        ip = sock.getsockname()[0]
        if ip and not ip.startswith('127.'):
            return ip
    except Exception:
        pass
    finally:
        if sock is not None:
            try:
                sock.close()
            except Exception:
                pass

    try:
        fallback = socket.gethostbyname(socket.gethostname())
        if fallback and not fallback.startswith('127.'):
            return fallback
    except Exception:
        pass

    return ''


def get_local_network_url(request_obj):
    """Return a QR base URL that works for devices on the same network."""
    # Optional override for reverse proxies or custom domains.
    configured = (os.environ.get('QR_BASE_URL') or '').strip().rstrip('/')
    if configured:
        return configured

    base = request_obj.host_url.rstrip('/')
    parsed = urlparse(base)
    host = (parsed.hostname or '').strip().lower()

    # Replace local-only hosts with LAN IP so phones/tablets can open QR links.
    if host in ('localhost', '127.0.0.1', '::1', '0.0.0.0'):
        lan_ip = _detect_lan_ip()
        if lan_ip:
            port = f":{parsed.port}" if parsed.port else ''
            netloc = f"{lan_ip}{port}"
            return urlunparse((parsed.scheme, netloc, '', '', '', '')).rstrip('/')

    return base


@app.route('/api/generate-qr-codes', methods=['POST'])
@login_required
def api_generate_qr_codes():
    """Generate QR codes for tables with links to the chatbot."""
    try:
        restaurant_id = get_current_restaurant_id()
        cfg = load_config(restaurant_id)
        
        data = request.get_json() or {}
        count = data.get('count', 10)
        start_table = data.get('start_table', 1)
        
        # Validate inputs
        if not isinstance(count, int) or count < 1 or count > 500:
            return jsonify({'error': 'Count must be between 1 and 500'}), 400
        if not isinstance(start_table, int) or start_table < 1:
            return jsonify({'error': 'Starting table number must be at least 1'}), 400
        
        qr_codes = []
        establishment_name = cfg.get('establishment_name', 'Restaurant')
        
        # Get the appropriate base URL for QR codes
        base_url = get_local_network_url(request)
        
        for i in range(count):
            table_number = start_table + i
            
            # Build the chatbot URL with table info using local network IP if applicable
            chatbot_url = f"{base_url}/chatbot?restaurant_id={restaurant_id}&table={table_number}"
            
            # Generate QR code
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=2,
            )
            qr.add_data(chatbot_url)
            qr.make(fit=True)
            
            # Create image
            img = qr.make_image(fill_color="black", back_color="white")
            
            # Convert to base64
            buffer = BytesIO()
            img.save(buffer, format='PNG')
            buffer.seek(0)
            img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
            image_url = f"data:image/png;base64,{img_base64}"
            
            qr_codes.append({
                'table_number': table_number,
                'image': image_url,
                'url': chatbot_url
            })
        
        return jsonify({
            'count': count,
            'qr_codes': qr_codes,
            'establishment': establishment_name
        })
    
    except Exception as e:
        app.logger.exception('QR code generation failed')
        return jsonify({'error': f'Failed to generate QR codes: {str(e)}'}), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
