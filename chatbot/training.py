"""
Training Data Management Module
================================
Handles loading, processing, and retrieval of custom training data for
restaurant-specific AI knowledge. Implements semantic search based on
token matching for context-aware responses.

Core Components:
  - Training data storage and manifest management
  - File handling for TXT, JSON, CSV formats
  - Text chunking with overlap for context preservation
  - Token-based semantic search
  - Context retrieval for chatbot prompts

Key Functions:
  - get_training_dir(): Get restaurant-specific training directory
  - get_training_manifest_path(): Access training manifest
  - load_training_manifest(): Load file metadata from manifest.json
  - build_training_context(): Retrieve relevant training chunks for queries
  - _read_text_file(): Safe file reading with error handling
  - _normalize_text(): Clean and normalize text data
  - _chunk_text(): Split text into overlapping chunks for context
  - _tokenize(): Extract searchable tokens from queries

Features:
  - Multi-restaurant isolation (restaurant_id based)
  - Semantic search scoring based on token frequency
  - Configurable chunk size (default 800 chars) with overlap (100 chars)
  - Manifest-based file tracking with original names
  - Safe file operations with UTF-8 handling
  - Query-based chunk scoring and ranking
  - Configurable result limiting (top N chunks)

Data Structure:
  training_data/
  ├── {restaurant_id}/
  │   ├── manifest.json (file metadata)
  │   ├── {uuid}.txt (training content)
  │   ├── {uuid}.json
  │   └── {uuid}.csv

Manifest Entry Format:
  {
    "id": "file_uuid",
    "original_name": "filename.txt",
    "stored_name": "hash.txt",
    "uploaded_at": "ISO timestamp",
    "size_bytes": 1024
  }
"""

import io
import json
import re
from pathlib import Path
from typing import Any

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

TRAINING_DIR = Path(__file__).resolve().parent.parent / 'training_data'
TRAINING_DIR.mkdir(parents=True, exist_ok=True)

TEXT_EXTENSIONS = {'.txt', '.json', '.csv', '.pdf', '.docx'}
SLIDING_CHUNK_SIZE = 600
SLIDING_OVERLAP_RATIO = 0.10


def get_training_dir(restaurant_id: str):
    if restaurant_id:
        safe_id = str(restaurant_id)
    else:
        safe_id = 'default'
    path = TRAINING_DIR / safe_id
    path.mkdir(parents=True, exist_ok=True)
    return path


def get_training_manifest_path(restaurant_id: str):
    return get_training_dir(restaurant_id) / 'manifest.json'


def get_training_history_path(restaurant_id: str):
    return get_training_dir(restaurant_id) / 'history.json'


def load_training_manifest(restaurant_id: str):
    manifest_path = get_training_manifest_path(restaurant_id)
    if manifest_path.exists():
        try:
            return json.loads(manifest_path.read_text())
        except Exception:
            return []
    return []


def save_training_manifest(restaurant_id: str, entries):
    manifest_path = get_training_manifest_path(restaurant_id)
    manifest_path.write_text(json.dumps(entries, indent=2, default=str))


def load_training_history(restaurant_id: str):
    history_path = get_training_history_path(restaurant_id)
    if history_path.exists():
        try:
            return json.loads(history_path.read_text())
        except Exception:
            return []
    return []


def save_training_history(restaurant_id: str, entries):
    history_path = get_training_history_path(restaurant_id)
    history_path.write_text(json.dumps(entries, indent=2, default=str))


def _read_text_file(path: Path):
    try:
        return path.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return ''


def _read_pdf_text(path: Path):
    if not PdfReader:
        return ''
    try:
        data = path.read_bytes()
        reader = PdfReader(io.BytesIO(data))
    except Exception:
        return ''
    parts = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ''
        except Exception:
            text = ''
        if text:
            parts.append(text)
    return '\n'.join(parts)


def _read_pdf_pages(path: Path):
    if not PdfReader:
        return []
    try:
        data = path.read_bytes()
        reader = PdfReader(io.BytesIO(data))
    except Exception:
        return []

    pages = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ''
        except Exception:
            text = ''
        text = text.strip()
        if text:
            pages.append((index, text))
    return pages


def _read_docx_text(path: Path):
    if not Document:
        return ''
    try:
        doc = Document(str(path))
    except Exception:
        return ''
    parts = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or '').strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


def _read_training_text(path: Path):
    suffix = path.suffix.lower()
    if suffix == '.pdf':
        return _read_pdf_text(path)
    if suffix == '.docx':
        return _read_docx_text(path)
    return _read_text_file(path)


def _normalize_text(text: str):
    return re.sub(r"\s+", " ", text or "").strip()


def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 100):
    if not text:
        return []
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = min(length, start + chunk_size)
        chunks.append(text[start:end])
        if end == length:
            break
        start = max(0, end - overlap)
    return chunks


def _split_non_empty_lines(text: str):
    return [line.strip() for line in (text or '').splitlines() if line.strip()]


def _is_heading_line(line: str):
    if not line:
        return False
    stripped = line.strip()
    if len(stripped) > 100:
        return False

    upperish = re.sub(r"[^A-Z]", "", stripped)
    letters = re.sub(r"[^A-Za-z]", "", stripped)
    if letters and (len(upperish) / len(letters)) >= 0.80 and len(stripped.split()) <= 10:
        return True

    if re.match(r"^(section|article|chapter)\s+\d+(?:\.\d+)*\b", stripped, re.IGNORECASE):
        return True
    if re.match(r"^\d+(?:\.\d+)+\s+", stripped):
        return True
    if re.match(r"^[A-Z][A-Za-z0-9\s\-/]{2,80}:$", stripped):
        return True
    return False


def _extract_visible_identifier(text: str):
    if not text:
        return None

    match = re.search(r"\b(invoice|inv)\s*#?:?\s*([A-Za-z0-9\-/]+)", text, re.IGNORECASE)
    if match:
        return f"Invoice {match.group(2)}"

    match = re.search(r"\b(section)\s+(\d+(?:\.\d+)*)\b", text, re.IGNORECASE)
    if match:
        return f"Section {match.group(2)}"

    match = re.search(r"\b(order)\s*#?:?\s*([A-Za-z0-9\-/]+)", text, re.IGNORECASE)
    if match:
        return f"Order {match.group(2)}"

    return None


def _sliding_chunks(text: str, chunk_size: int = SLIDING_CHUNK_SIZE, overlap_ratio: float = SLIDING_OVERLAP_RATIO):
    normalized = _normalize_text(text)
    if not normalized:
        return []

    overlap = max(1, int(chunk_size * overlap_ratio))
    step = max(1, chunk_size - overlap)
    chunks = []
    start = 0
    length = len(normalized)

    while start < length:
        end = min(length, start + chunk_size)
        chunk_text = normalized[start:end].strip()
        if chunk_text:
            chunks.append(chunk_text)
        if end >= length:
            break
        start += step
    return chunks


def _docx_structured_sections(path: Path):
    if not Document:
        return []
    try:
        doc = Document(str(path))
    except Exception:
        return []

    sections = []
    current_title = None
    current_lines = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or '').strip()
        if not text:
            continue

        style_name = (getattr(paragraph.style, 'name', '') or '').lower()
        looks_heading = ('heading' in style_name) or _is_heading_line(text)

        if looks_heading:
            if current_lines:
                sections.append({
                    'section_title': current_title,
                    'content': '\n'.join(current_lines),
                })
                current_lines = []
            current_title = text
            continue

        current_lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                current_lines.append(' | '.join(cells))

    if current_lines:
        sections.append({
            'section_title': current_title,
            'content': '\n'.join(current_lines),
        })

    return sections


def _pdf_structured_sections(path: Path):
    pages = _read_pdf_pages(path)
    if not pages:
        return []

    sections = []
    for page_number, page_text in pages:
        lines = _split_non_empty_lines(page_text)
        if not lines:
            continue

        current_title = None
        current_lines = []

        for line in lines:
            if _is_heading_line(line):
                if current_lines:
                    sections.append({
                        'page': page_number,
                        'section_title': current_title,
                        'content': '\n'.join(current_lines),
                    })
                    current_lines = []
                current_title = line
                continue
            current_lines.append(line)

        if current_lines:
            sections.append({
                'page': page_number,
                'section_title': current_title,
                'content': '\n'.join(current_lines),
            })

    return sections


def _is_structured_document(file_ext: str, raw_text: str):
    if file_ext in {'.pdf', '.docx'}:
        return True
    lines = _split_non_empty_lines(raw_text)
    if not lines:
        return False
    heading_like = sum(1 for line in lines[:80] if _is_heading_line(line))
    return heading_like >= 2


def _make_chunk(content: str, metadata: dict[str, Any]):
    chunk_content = _normalize_text(content)
    if not chunk_content:
        return None

    identifier = _extract_visible_identifier(content)
    chunk_meta = dict(metadata)
    if identifier:
        chunk_meta['identifier'] = identifier

    return {
        'content': chunk_content,
        'metadata': chunk_meta,
    }


def build_training_chunks(restaurant_id: str, file_path: Path, entry: dict):
    file_ext = file_path.suffix.lower()
    original_name = entry.get('original_name') or file_path.name

    raw_text = _read_training_text(file_path)
    if not raw_text:
        return []

    base_meta = {
        'restaurant_id': restaurant_id,
        'file_ext': file_ext,
        'source_file': original_name,
    }

    chunks = []
    if _is_structured_document(file_ext, raw_text):
        sections = []
        if file_ext == '.pdf':
            sections = _pdf_structured_sections(file_path)
        elif file_ext == '.docx':
            sections = _docx_structured_sections(file_path)

        for section in sections:
            metadata = dict(base_meta)
            if section.get('page') is not None:
                metadata['page'] = section.get('page')
            if section.get('section_title'):
                metadata['section_title'] = section.get('section_title')
            chunk = _make_chunk(section.get('content', ''), metadata)
            if chunk:
                chunks.append(chunk)

    if not chunks:
        for chunk_text in _sliding_chunks(raw_text):
            metadata = dict(base_meta)
            chunk = _make_chunk(chunk_text, metadata)
            if chunk:
                chunks.append(chunk)

    return chunks


def _tokenize(query: str):
    tokens = re.findall(r"[a-z0-9]+", (query or "").lower())
    return [t for t in tokens if len(t) > 2]


def build_training_context(restaurant_id: str, query: str, max_chunks: int = 3):
    tokens = _tokenize(query)
    if not tokens:
        return ''

    entries = load_training_manifest(restaurant_id)
    training_dir = get_training_dir(restaurant_id)

    scored_chunks = []
    for entry in entries:
        stored_name = entry.get('stored_name')
        original_name = entry.get('original_name') or stored_name
        if not stored_name:
            continue
        file_path = training_dir / stored_name
        if not file_path.exists():
            continue
        if file_path.suffix.lower() not in TEXT_EXTENSIONS:
            continue
        file_chunks = build_training_chunks(restaurant_id, file_path, entry)
        if not file_chunks:
            continue
        for chunk_obj in file_chunks:
            chunk = chunk_obj.get('content', '')
            metadata = chunk_obj.get('metadata') or {}
            score = sum(chunk.lower().count(t) for t in tokens)
            if score <= 0:
                continue
            scored_chunks.append((score, original_name, chunk, metadata))

    if not scored_chunks:
        return ''

    scored_chunks.sort(key=lambda item: item[0], reverse=True)
    picked = scored_chunks[:max_chunks]
    blocks = []
    for _, name, chunk, metadata in picked:
        context_header = [f"File: {name}"]
        if metadata.get('page') is not None:
            context_header.append(f"Page: {metadata['page']}")
        if metadata.get('section_title'):
            context_header.append(f"Section: {metadata['section_title']}")
        if metadata.get('identifier'):
            context_header.append(f"Identifier: {metadata['identifier']}")
        blocks.append(f"{' | '.join(context_header)}\n{chunk}")
    return "\n\n".join(blocks)
